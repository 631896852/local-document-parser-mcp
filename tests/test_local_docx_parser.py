import tempfile
import unittest
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import fitz
from openpyxl import Workbook

from src.local_docx_parser import (
    LocalDocxParseService,
    _extract_markdown_tables,
    _merge_ocr_heading_lines,
)


def build_sample_docx(path: Path) -> None:
    doc = Document()
    doc.add_paragraph("1.总论")
    doc.add_paragraph("项目背景包含投资估算。")
    doc.add_paragraph("1.1范围")
    doc.add_paragraph("范围正文。")
    doc.add_paragraph("2.经济评价")
    doc.add_paragraph("表2-1 项目总投资估算表")
    table = doc.add_table(rows=3, cols=3)
    table.cell(0, 0).text = "序号"
    table.cell(0, 1).text = "项目"
    table.cell(0, 2).text = "金额"
    table.cell(1, 0).text = "1"
    table.cell(1, 1).text = "建设投资"
    table.cell(1, 2).text = "  100"
    table.cell(2, 0).text = "2"
    table.cell(2, 1).text = "运维投资"
    table.cell(2, 2).text = "20"
    doc.save(path)


def build_automatic_heading_docx(path: Path) -> None:
    doc = Document()

    def add_heading(text: str, level: int) -> None:
        paragraph = doc.add_paragraph(text, style=f"Heading {level}")
        properties = paragraph._p.get_or_add_pPr()
        num_properties = OxmlElement("w:numPr")
        list_level = OxmlElement("w:ilvl")
        list_level.set(qn("w:val"), str(level - 1))
        num_id = OxmlElement("w:numId")
        num_id.set(qn("w:val"), "2")
        num_properties.append(list_level)
        num_properties.append(num_id)
        properties.append(num_properties)

    add_heading("设计依据", 1)
    add_heading("依据的批准文件", 2)
    doc.add_paragraph("批准文件正文。")
    add_heading("执行的主要标准规范", 2)
    doc.add_paragraph("标准规范正文。")
    add_heading("工程概况", 1)
    add_heading("建设项目概况", 2)
    add_heading("建设规模和主要工作量", 3)
    doc.add_paragraph("建设规模正文。")
    doc.save(path)


def build_sample_pdf(path: Path) -> None:
    doc = fitz.open()
    page = doc.new_page()
    lines = [
        "1",
        "Contents",
        "1 Summary............................. 1",
        "2 Economy............................. 2",
        "1 Summary",
        "1.1 Background",
        "Background includes investment estimate.",
        "2 Economy",
        "2.1 Investment estimate",
        "Table 2-1 Project investment table",
        "No Item Amount",
        "1 Construction 100",
        "2 Operation 20",
    ]
    y = 72
    for line in lines:
        page.insert_text((72, y), line, fontname="china-s")
        y += 18
    doc.save(path)
    doc.close()


def build_grid_table_pdf(path: Path) -> None:
    doc = fitz.open()
    page = doc.new_page()
    page.insert_text((72, 50), "1 Summary", fontname="china-s")
    page.insert_text((72, 68), "Background.", fontname="china-s")
    page.insert_text((72, 92), "2 Economy", fontname="china-s")
    page.insert_text((72, 117), "Table 2-1 Project investment table", fontname="china-s")
    x0, y0 = 72, 142
    cols = [0, 90, 250, 360]
    rows = [0, 32, 64, 96]
    for x in cols:
        page.draw_line((x0 + x, y0), (x0 + x, y0 + rows[-1]))
    for y in rows:
        page.draw_line((x0, y0 + y), (x0 + cols[-1], y0 + y))
    values = [
        ["No", "Item", "Amount"],
        ["1", "Construction", "100"],
        ["2", "Operation", "20"],
    ]
    for row_index, row in enumerate(values):
        for col_index, value in enumerate(row):
            page.insert_text(
                (x0 + cols[col_index] + 8, y0 + rows[row_index] + 21),
                value,
                fontname="china-s",
            )
    doc.save(path)
    doc.close()


def build_scanned_pdf(path: Path) -> None:
    doc = fitz.open()
    page = doc.new_page()
    pixmap = fitz.Pixmap(fitz.csRGB, fitz.IRect(0, 0, 600, 800), False)
    pixmap.clear_with(255)
    page.insert_image(page.rect, stream=pixmap.tobytes("png"))
    doc.save(path)
    doc.close()


def build_chinese_fee_pdf(path: Path) -> None:
    doc = fitz.open()
    page = doc.new_page()
    lines = [
        "目录",
        "第一章 总则 ........................................ - 1 -",
        "第二章 其他费用 .................................... - 4 -",
        "- 1 -",
        "第一章 总则",
        "一、总则正文包含其他费用。",
        "表1 工程建设其他费用和预备费组成表",
        "费用项目类别 费用项目名称",
        "其他费用",
        "1.可行性研究费",
        "2.项目核准申请报告编制费",
        "2 油气田工程（评价面积S，单位km2）",
        "预备费",
        "45.基本预备费",
        "3 城市规划区外林地（属于经营性建设项目）",
        "- 4 -",
        "第二章 其他费用",
        "3. 土地补偿费",
        "(1) 费用内容",
        "土地补偿费是指建设项目使用土地应支付的费用，包括征地补偿费和临时",
        "土地补偿费，以及由于使用土地发生的其他有关费用。",
        "24. 设计审查费",
        "(1) 费用内容 设计审查费包括施工图审查和初步设计审查两项内容。",
        "(2) 计算方法 初步设计审查费＝基本设计费×6％。",
        "45. 基本预备费",
        "基本预备费＝（工程费用＋其他费用）×基本预备费费率",
    ]
    y = 72
    for line in lines:
        page.insert_text((72, y), line, fontname="china-s")
        y += 18
    doc.save(path)
    doc.close()


def build_chinese_marker_pdf(path: Path) -> None:
    doc = fitz.open()
    page = doc.new_page()
    lines = [
        "一、费用范围",
        "费用范围正文第一行",
        "费用范围正文第二行",
        "二、计算方法",
        "计算方法正文",
        "（一）基本公式",
        "基本公式正文",
        "（二）调整系数",
        "调整系数正文",
        "三、计列要求",
        "计列要求正文",
    ]
    y = 72
    for line in lines:
        page.insert_text((72, y), line, fontname="china-s")
        y += 18
    doc.save(path)
    doc.close()


def build_wrapped_chinese_project_pdf(path: Path) -> None:
    doc = fitz.open()
    page = doc.new_page()
    lines = [
        "一、示例项目建设工程",
        "1. 工程主要内容：部署新井。",
        "二、车辆管理中心2025年车辆更新项目",
        "1. 工程主要内容：更新车辆。",
        "九、地面工程维修中心2025年设备更新项目",
        "1. 工程主要内容：更新设备。",
        "十、应急救援中心2025年消防车辆更新项目 1. 工程主要内容：更新16吨泡沫消防车8台。",
        "2. 工程投资控制在2197万元以内。",
        "十一、油区护卫管理中心2025年巡检装备购置项目",
        "1. 工程主要内容：购置无人机。",
    ]
    y = 72
    for line in lines:
        page.insert_text((72, y), line, fontname="china-s")
        y += 18
    doc.save(path)
    doc.close()


def build_sample_xlsx(path: Path) -> None:
    wb = Workbook()
    ws = wb.active
    ws.title = "项目基本信息"
    ws.merge_cells("A1:C1")
    ws["A1"] = "项目基本信息表"
    ws.append([1, "项目名称", ""])
    ws.append([2, "建设单位", "滨南采油厂"])
    ws.append([3, "新老区", "老区"])
    ws2 = wb.create_sheet("经济评价（单价维护）")
    ws2.append(["时间", "项目名称", "费用"])
    ws2.append([46113, "伴生气价格", ""])
    ws2.append([46114, "燃气单价", ""])
    ws3 = wb.create_sheet("总收入")
    ws3.append(["时间", "股票期末余额\n（港币）", "股币总额\n（人民币）"])
    ws3.append([46082, 369518, 450121.2378])
    ws3.append([46113, 445600, 526585.56])
    wb.save(path)


class LocalDocxParseServiceTest(unittest.TestCase):
    def setUp(self):
        self.tmp = tempfile.TemporaryDirectory()
        self.path = Path(self.tmp.name) / "sample.docx"
        build_sample_docx(self.path)
        self.service = LocalDocxParseService(self.path, file_id="sample-id")

    def tearDown(self):
        self.tmp.cleanup()

    def test_summary_contains_headings(self):
        text = self.service.extract_file_summary_text()
        self.assertIn("<file-id>sample-id</file-id>", text)
        self.assertIn("1.总论", text)
        self.assertIn("2.经济评价", text)

    def test_rebuilds_automatic_heading_numbers_after_doc_conversion(self):
        path = Path(self.tmp.name) / "automatic-headings.docx"
        build_automatic_heading_docx(path)
        service = LocalDocxParseService(path, file_id="automatic-headings")

        headings = [chapter.full_title for chapter in service.chapters]
        self.assertIn("1 设计依据", headings)
        self.assertIn("1.1 依据的批准文件", headings)
        self.assertIn("1.2 执行的主要标准规范", headings)
        self.assertIn("2 工程概况", headings)
        self.assertIn("2.1.1 建设规模和主要工作量", headings)

    def test_chapter_matching(self):
        by_number = self.service.extract_chapter_summary("1")
        by_name = self.service.extract_chapter_summary("总论")
        self.assertIn("1.总论", by_number)
        self.assertIn("1.总论", by_name)

    def test_read_lines_clips_bounds(self):
        text = self.service.read_lines(0, 2)
        self.assertIn("1.总论", text)

    def test_search_uses_or_semantics(self):
        text = self.service.search_content("投资估算,范围", page_no=1, page_size=10)
        self.assertIn("搜索结果总条数:4", text)

    def test_table_list_and_content(self):
        titles = self.service.extract_table_list("2")
        self.assertIn("表2-1 项目总投资估算表", titles)
        content = self.service.extract_table_content("表2-1 项目总投资估算表", row_scope="1-2")
        self.assertIn("总行数:3", content)
        self.assertIn("|序号|项目|金额|", content)
        self.assertIn("|1|建设投资|  100|", content)

    def test_table_content_falls_back_to_keyword_search(self):
        content = self.service.extract_table_content("不存在的表", row_keywords="投资估算")
        self.assertIn("未找到匹配表格标题", content)
        self.assertIn("项目背景包含投资估算", content)


class LocalPdfParseServiceTest(unittest.TestCase):
    def setUp(self):
        self.tmp = tempfile.TemporaryDirectory()
        self.path = Path(self.tmp.name) / "sample.pdf"
        build_sample_pdf(self.path)
        self.service = LocalDocxParseService(self.path, file_id="pdf-id")

    def tearDown(self):
        self.tmp.cleanup()

    def test_pdf_summary_and_chapters(self):
        text = self.service.extract_file_summary_text()
        self.assertIn("<file-id>pdf-id</file-id>", text)
        self.assertIn("解析引擎:pdf-inspector + PyMuPDF", text)
        self.assertIn("PDF类型:text_based", text)
        self.assertIn("1 Summary", text)
        self.assertIn("2 Economy", text)

    def test_ocr_heading_lines_are_merged_without_merging_table_rows(self):
        lines = _merge_ocr_heading_lines(
            [
                "1范围",
                "2",
                "规范性引用文件",
                "3. 11",
                "注水 water injection",
                "1",
                "站长",
            ]
        )
        self.assertIn("2 规范性引用文件", lines)
        self.assertIn("3.11 注水 waterinjection", lines)
        self.assertIn("1", lines)
        self.assertIn("站长", lines)

    def test_ocr_toc_page_is_not_merged_into_body_headings(self):
        lines = _merge_ocr_heading_lines(["目", "次", "1", "范围", "2", "规范性引用文件"])
        self.assertEqual(lines, ["目录", "1", "范围", "2", "规范性引用文件"])

    def test_pdf_inspection_and_markdown_are_available(self):
        inspection = self.service.inspect_pdf(include_images=False)
        self.assertIn('"pdfType": "text_based"', inspection)
        self.assertIn('"hasMarkdown": true', inspection)

        markdown = self.service.extract_pdf_markdown()
        self.assertIn("Contents", markdown)
        self.assertIn("Project investment table", markdown)

    def test_pdf_inspector_markdown_tables_keep_caption_and_cells(self):
        tables = _extract_markdown_tables(
            "表2-1 项目投资表\n\n|项目|金额|\n|---|---:|\n|建设投资|100|"
        )
        self.assertEqual(len(tables), 1)
        self.assertEqual(tables[0][0], "表2-1 项目投资表")
        self.assertEqual(tables[0][2][1], ["建设投资", "100"])

    def test_scanned_pdf_reports_ocr_page_and_image_region(self):
        path = Path(self.tmp.name) / "scanned.pdf"
        build_scanned_pdf(path)
        service = LocalDocxParseService(path, file_id="scanned-pdf")

        inspection = service.inspect_pdf(include_images=True)
        self.assertIn('"pdfType": "scanned"', inspection)
        self.assertIn('"pagesNeedingOcr": [\n    1\n  ]', inspection)
        self.assertIn('"name": "fzImg0"', inspection)

        markdown = service.extract_pdf_markdown()
        self.assertIn("需要 OCR 的页面: 1", markdown)

    def test_pdf_search_and_table_fallback(self):
        search = self.service.search_content("investment", page_no=1, page_size=10)
        self.assertIn("搜索结果总条数:", search)
        titles = self.service.extract_table_list("2")
        self.assertIn("Table 2-1", titles)
        content = self.service.extract_table_content("Table 2-1 Projectinvestmenttable")
        self.assertIn("Construction", content)

    def test_pdf_detected_table_keeps_grid_structure(self):
        path = Path(self.tmp.name) / "grid-table.pdf"
        build_grid_table_pdf(path)
        service = LocalDocxParseService(path, file_id="grid-pdf")

        titles = service.extract_table_list("2")
        self.assertIn("Table 2-1 Projectinvestmenttable", titles)

        content = service.extract_table_content(
            "Table 2-1 Projectinvestmenttable",
            col_keywords="Amount",
        )
        self.assertIn("|Amount|", content)
        self.assertIn("|100|", content)
        self.assertIn("|20|", content)
        self.assertNotIn("|Construction|", content)

    def test_pdf_chinese_fee_headings(self):
        path = Path(self.tmp.name) / "fees.pdf"
        build_chinese_fee_pdf(path)
        service = LocalDocxParseService(path, file_id="fee-pdf")
        summary = service.extract_file_summary_text()
        self.assertIn("第一章 总则", summary)
        self.assertIn("第二章 其他费用", summary)
        self.assertIn("3. 土地补偿费", summary)
        chapter = service.extract_chapter_content("土地补偿费")
        self.assertIn("征地补偿费和临时 土地补偿费", chapter)
        self.assertNotIn("- 4 -", chapter)
        self.assertNotIn("1.可行性研究费", summary)
        self.assertNotIn("45.基本预备费", summary)

    def test_pdf_summary_filters_table_like_numbered_headings(self):
        path = Path(self.tmp.name) / "fees.pdf"
        build_chinese_fee_pdf(path)
        service = LocalDocxParseService(path, file_id="fee-pdf")
        summary = service.extract_file_summary_text()
        self.assertIn("第二章 其他费用", summary)
        self.assertIn("3. 土地补偿费", summary)
        self.assertNotIn("2 油气田工程", summary)
        self.assertNotIn("城市规划区外林地", summary)

    def test_pdf_chinese_marker_headings_without_explicit_chapters(self):
        path = Path(self.tmp.name) / "markers.pdf"
        build_chinese_marker_pdf(path)
        service = LocalDocxParseService(path, file_id="marker-pdf")
        summary = service.extract_file_summary_text()
        self.assertIn("一、费用范围", summary)
        self.assertIn("二、计算方法", summary)
        self.assertIn("（一）基本公式", summary)
        self.assertIn("（二）调整系数", summary)
        chapter = service.extract_chapter_content("计算方法")
        self.assertIn("计算方法正文", chapter)
        self.assertIn("基本公式正文", chapter)

    def test_pdf_splits_chinese_project_heading_joined_with_body(self):
        path = Path(self.tmp.name) / "wrapped-projects.pdf"
        build_wrapped_chinese_project_pdf(path)
        service = LocalDocxParseService(path, file_id="wrapped-projects")
        summary = service.extract_file_summary_text()
        self.assertIn("十、应急救援中心2025年消防车辆更新项目", summary)
        chapter = service.extract_chapter_content("应急救援中心")
        self.assertIn("更新16吨泡沫消防车8台", chapter)
        self.assertNotIn("十一、油区护卫管理中心", chapter)

    def test_pdf_compound_chapter_query(self):
        path = Path(self.tmp.name) / "fees.pdf"
        build_chinese_fee_pdf(path)
        service = LocalDocxParseService(path, file_id="fee-pdf")
        summary = service.extract_chapter_summary("第二章 3")
        self.assertIn("3. 土地补偿费", summary)

    def test_pdf_query_content_handles_natural_language_question(self):
        path = Path(self.tmp.name) / "fees.pdf"
        build_chinese_fee_pdf(path)
        service = LocalDocxParseService(path, file_id="fee-pdf")
        result = service.query_content("审查设计文件要收什么钱", page_no=1, page_size=3)
        self.assertIn("24. 设计审查费", result)
        self.assertIn("施工图审查", result)

    def test_pdf_summary_limits_headings_to_third_level(self):
        path = Path(self.tmp.name) / "policy.pdf"
        doc = fitz.open()
        page = doc.new_page()
        lines = [
            "1 总则",
            "总则正文",
            "2 术语和定义",
            "术语正文",
            "3 投资项目权限划分和决策主体",
            "权限正文",
            "4 组织管理与职责",
            "职责正文",
            "5 管理程序及要求",
            "5.2 项目可行性研究",
            "5.2.2 可行性研究报告评估论证",
            "5.2.2.1 境内独资项目可行性研究报告评估论证",
            "正文内容",
            "5.5 项目实施管理",
            "5.5.2 设计变更",
            "5.5.2.3 重大设计变更的界定",
            "正文内容",
        ]
        y = 72
        for line in lines:
            page.insert_text((72, y), line, fontname="china-s")
            y += 18
        doc.save(path)
        doc.close()
        service = LocalDocxParseService(path, file_id="policy-pdf")
        summary = service.extract_file_summary_text()
        self.assertIn("5.2.2 可行性研究报告评估论证", summary)
        self.assertIn("5.5.2 设计变更", summary)
        self.assertNotIn("5.2.2.1 境内独资项目可行性研究报告评估论证", summary)
        self.assertNotIn("5.5.2.3 重大设计变更的界定", summary)


class LocalExcelParseServiceTest(unittest.TestCase):
    def setUp(self):
        self.tmp = tempfile.TemporaryDirectory()
        self.path = Path(self.tmp.name) / "sample.xlsx"
        build_sample_xlsx(self.path)
        self.service = LocalDocxParseService(self.path, file_id="excel-id")

    def tearDown(self):
        self.tmp.cleanup()

    def test_excel_summary_lists_sheets(self):
        text = self.service.extract_file_summary_text()
        self.assertIn("<file-id>excel-id</file-id>", text)
        self.assertIn('<sheet name="项目基本信息" contentType="html">', text)
        self.assertIn('<sheet name="经济评价（单价维护）"', text)

    def test_excel_sheet_summary_and_content(self):
        summary = self.service.extract_chapter_summary("项目基本信息")
        self.assertIn("字符总数:", summary)
        self.assertIn("# 前两列预览", summary)
        self.assertIn('colspan="3"', summary)
        content = self.service.extract_chapter_content("项目基本信息")
        self.assertIn("<table>", content)
        self.assertIn("滨南采油厂", content)
        by_index = self.service.extract_chapter_content("2")
        self.assertIn('sheet="经济评价（单价维护）" row="2"', by_index)
        self.assertIn('cell="A2" header="时间">2026-04-01（4月）', by_index)
        self.assertIn("2026-04-01（4月）", by_index)
        self.assertIn("伴生气价格", by_index)

    def test_excel_read_lines_and_search(self):
        lines = self.service.read_lines(1, 2)
        self.assertIn('<sheet name="项目基本信息"', lines)
        self.assertIn("<table>", lines)
        search = self.service.search_content("滨南采油厂", page_no=1, page_size=10)
        self.assertIn("搜索结果总条数:1", search)
        self.assertIn("滨南采油厂", search)
        month_search = self.service.search_content("4月", page_no=1, page_size=10)
        self.assertIn("2026-04-01（4月）", month_search)

    def test_excel_table_tools_match_remote_limit(self):
        titles = self.service.extract_table_list("项目基本信息")
        self.assertIn('sheet name="项目基本信息"', titles)
        content = self.service.extract_table_content("项目基本信息表")
        self.assertIn("项目基本信息!C3", content)
        self.assertIn("滨南采油厂", content)

    def test_excel_table_content_matches_row_and_column_keywords(self):
        content = self.service.extract_table_content(
            "总收入",
            row_keywords="4月",
            col_keywords="股票期末余额",
        )
        self.assertIn("总收入!B3", content)
        self.assertIn('header="股票期末余额 （港币）"', content)
        self.assertIn("445600", content)

    def test_excel_table_content_can_treat_row_keyword_as_header(self):
        content = self.service.extract_table_content("总收入", row_keywords="股票期末余额")
        self.assertIn("总收入!B3", content)
        self.assertIn("445600", content)


if __name__ == "__main__":
    unittest.main()
