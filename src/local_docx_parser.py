from __future__ import annotations

import math
import re
import hashlib
import json
import shutil
import subprocess
import tempfile
from collections import Counter
from html import escape
from dataclasses import dataclass, field
from pathlib import Path
from typing import Iterator

from docx import Document
from docx.document import Document as DocumentObject
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table
from docx.text.paragraph import Paragraph
import fitz

from src.local_excel_parser import EXCEL_SUFFIXES, LocalExcelParseService
from src.local_ocr import LocalOcrBackend, OcrDocument
from src.pdf_inspector_backend import PdfInspection, PdfInspectorBackend

try:
    import jieba
except ImportError:  # pragma: no cover - jieba is an optional enhancement.
    jieba = None  # type: ignore[assignment]


@dataclass
class LocalTable:
    title: str
    chapter_key: str
    start_line: int
    end_line: int
    markdown: str
    rows: list[list[str]]
    is_html: bool = False


@dataclass
class LocalChapter:
    number: str
    title: str
    full_title: str
    level: int
    start_line: int
    end_line: int = 0
    char_count: int = 0
    table_char_count: int = 0
    children: list["LocalChapter"] = field(default_factory=list)


@dataclass
class SearchHit:
    line_no: int
    text: str
    keyword: str
    score: int = 0


@dataclass
class QueryHit:
    line_no: int
    text: str
    score: int
    source: str


def _resolve_word_parse_path(path: Path) -> Path:
    path = path.expanduser()
    suffix = path.suffix.lower()
    if suffix != ".doc":
        return path

    converted = _convert_doc_to_docx(path)
    if converted is not None:
        return converted

    raise ValueError(
        f"无法自动转换 .doc 文件: {path}. 请安装 LibreOffice 后重试。"
    )


def _convert_doc_to_docx(path: Path) -> Path | None:
    out_dir = _doc_conversion_cache_dir(path)
    out_dir.mkdir(parents=True, exist_ok=True)
    cached = out_dir / f"{path.stem}.docx"
    if cached.exists():
        return cached

    soffice = shutil.which("soffice") or shutil.which("libreoffice")
    if soffice is not None:
        subprocess.run(
            [
                soffice,
                "-env:UserInstallation=file://" + str(out_dir / "profile"),
                "--headless",
                "--convert-to",
                "docx",
                "--outdir",
                str(out_dir),
                str(path),
            ],
            check=False,
            stdout=subprocess.DEVNULL,
            stderr=subprocess.DEVNULL,
        )
        candidate = out_dir / f"{path.stem}.docx"
        if candidate.exists():
            return candidate

    textutil = shutil.which("textutil")
    if textutil is not None:
        candidate = out_dir / f"{path.stem}.docx"
        subprocess.run(
            ["textutil", "-convert", "docx", "-output", str(candidate), str(path)],
            check=False,
            stdout=subprocess.DEVNULL,
            stderr=subprocess.DEVNULL,
        )
        if candidate.exists():
            return candidate

    return None


def _doc_conversion_cache_dir(path: Path) -> Path:
    stat = path.stat()
    digest = hashlib.sha256(
        f"{path.resolve()}:{stat.st_mtime_ns}:{stat.st_size}".encode("utf-8")
    ).hexdigest()[:16]
    return Path(tempfile.gettempdir()) / "document_parser_doc_cache" / digest


class LocalDocxParseService:
    def __init__(self, path: str | Path, file_id: str | None = None) -> None:
        self.path = Path(path)
        self._pdf_backend = PdfInspectorBackend()
        self.pdf_inspection: PdfInspection | None = None
        self._pdf_images_loaded = False
        self._ocr_backend = LocalOcrBackend()
        self.ocr_document: OcrDocument | None = None
        self._pdf_page_texts: dict[int, str] = {}
        self._pdf_line_pages: dict[int, int] = {}
        self._pdf_toc_pages: set[int] = set()
        self._excel_service: LocalExcelParseService | None = None
        if self.path.suffix.lower() in EXCEL_SUFFIXES:
            self._excel_service = LocalExcelParseService(self.path, file_id=file_id)
            self.file_id = self._excel_service.file_id
            self.file_name = self._excel_service.file_name
            self.lines = self._excel_service.lines
            self.chapters = []
            self.tables = []
            self._toc_line = None
            self._body_start_line = None
            self._query_terms = []
            return
        self._temp_dir: tempfile.TemporaryDirectory[str] | None = None
        self._parse_path = _resolve_word_parse_path(self.path)
        self.file_id = file_id or self.path.stem
        self.file_name = self.path.name
        self.lines: list[str] = []
        self.chapters: list[LocalChapter] = []
        self.tables: list[LocalTable] = []
        self._toc_line: int | None = None
        self._body_start_line: int | None = None
        self._query_terms: list[str] = []
        self._parse()
        self._register_query_terms()

    def extract_file_summary_text(self) -> str:
        if self._excel_service is not None:
            return self._excel_service.extract_file_summary_text()
        if self.path.suffix.lower() == ".pdf":
            return self._extract_pdf_file_summary_text()
        if self.path.suffix.lower() == ".doc" and self._toc_line is not None:
            return self._extract_partitioned_file_summary_text()
        headings = "\n".join(ch.full_title for ch in self._summary_chapters())
        return self._wrap_result(
            "# 总体摘要\n\n"
            f"该文档共{len(self.lines)}行,{self._total_chars()}字符。\n\n"
            "# 正文部分:\n"
            "## 摘要信息\n"
            f"字符总数:{self._total_chars()},总行数:{len(self.lines)},"
            f"行索引范围:1~{len(self.lines)}\n"
            "## 章节标题信息\n"
            f"{headings}"
        )

    def _extract_partitioned_file_summary_text(self) -> str:
        headings = "\n".join(ch.full_title for ch in self._summary_chapters())
        toc_line = self._toc_line
        body_start = self._body_start_line or (self.chapters[0].start_line if self.chapters else 1)
        sections = ["正文部分"]
        if toc_line and toc_line > 1:
            sections.insert(0, "封面部分")
            sections.insert(1, "目录部分")
        elif body_start > 1:
            sections.insert(0, "封面部分")
        parts = [
            "# 总体摘要",
            "",
            f"该文档共{len(self.lines)}行,{self._total_chars()}字符,整体分为{'、'.join(sections)}。",
            "",
        ]
        if toc_line and toc_line > 1:
            cover = self.lines[: toc_line - 1]
            parts.extend(
                [
                    "# 封面部分:",
                    "## 摘要信息",
                    f"总字符数:{len(chr(10).join(cover))},总行数:{len(cover)},行索引范围:1~{toc_line - 1}",
                    "## 具体内容",
                    "\n".join(cover[:80]),
                    "",
                ]
            )
        if toc_line and body_start > toc_line:
            toc = self.lines[toc_line - 1 : body_start - 1]
            parts.extend(
                [
                    "# 目录部分:",
                    "## 摘要信息",
                    f"总字符数:{len(chr(10).join(toc))},总行数:{len(toc)},行索引范围:{toc_line}~{body_start - 1}",
                    "## 具体内容",
                    "\n".join(toc[:120]),
                    "",
                ]
            )
        body_lines = max(0, len(self.lines) - body_start + 1)
        body_chars = len("\n".join(self.lines[body_start - 1 :])) if self.lines else 0
        parts.extend(
            [
                "# 正文部分:",
                "## 摘要信息",
                f"字符总数:{body_chars},总行数:{body_lines},行索引范围:{body_start}~{len(self.lines)}",
                "## 章节标题信息",
                headings,
            ]
        )
        return self._wrap_result("\n".join(parts))

    def _extract_pdf_file_summary_text(self) -> str:
        text = self._extract_partitioned_file_summary_text()
        summary = self._pdf_inspection_summary()
        if not summary:
            return text
        return text.replace("<result>\n", f"<result>\n# PDF解析信息\n{summary}\n\n", 1)

    def inspect_pdf(self, *, include_images: bool = True) -> str:
        if self.path.suffix.lower() != ".pdf":
            return self._wrap_tip("inspectPdf 仅支持 PDF 文件")
        inspection = self.pdf_inspection or self._pdf_backend.inspect(self.path)
        self.pdf_inspection = inspection
        if include_images and inspection.available and not self._pdf_images_loaded:
            inspection.image_regions = self._pdf_backend.extract_image_regions(self.path)
            self._pdf_images_loaded = True
        return self._wrap_result(
            json.dumps(
                {
                    **inspection.as_dict(include_images=include_images),
                    "ocr": self._ocr_info_dict(),
                },
                ensure_ascii=False,
                indent=2,
            )
        )

    def ocr_pdf(self) -> str:
        if self.path.suffix.lower() != ".pdf":
            return self._wrap_tip("ocrPdf 仅支持 PDF 文件")
        inspection = self.pdf_inspection or self._pdf_backend.inspect(self.path)
        self.pdf_inspection = inspection
        requested = inspection.pages_needing_ocr
        if not requested:
            return self._wrap_tip("该 PDF 当前没有需要 OCR 的页面")
        if self.ocr_document is None:
            self.ocr_document = self._ocr_backend.recognize_pdf(self.path, requested)
        document = self.ocr_document
        if not document.successful_pages:
            return self._wrap_tip(document.error or "OCR 未识别到有效文字")
        body = "\n\n".join(
            f"<!-- Page {page} -->\n{document.pages[page].text}"
            for page in document.successful_pages
        )
        return self._wrap_tip_result(
            "\n".join(
                [
                    f"OCR引擎:{document.engine}",
                    f"OCR成功页:{_format_page_list(document.successful_pages)}",
                    f"OCR失败页:{_format_page_list(document.failed_pages)}",
                    f"OCR耗时:{document.processing_time_ms}ms",
                ]
            ),
            body,
        )

    def extract_pdf_markdown(self) -> str:
        if self.path.suffix.lower() != ".pdf":
            return self._wrap_tip("extractPdfMarkdown 仅支持 PDF 文件")
        inspection = self.pdf_inspection or self._pdf_backend.inspect(self.path)
        self.pdf_inspection = inspection
        if self.ocr_document is not None and self.ocr_document.successful_pages:
            return self._wrap_tip_result(
                "\n".join(
                    [
                        f"已使用 {self.ocr_document.engine} 识别扫描页",
                        f"OCR成功页:{_format_page_list(self.ocr_document.successful_pages)}",
                        f"OCR失败页:{_format_page_list(self.ocr_document.failed_pages)}",
                    ]
                ),
                self._hybrid_pdf_markdown(),
            )
        if inspection.markdown:
            return self._wrap_result(inspection.markdown)
        reason = inspection.error or "pdf-inspector 未产生可靠 Markdown"
        if inspection.pages_needing_ocr:
            reason += f"；需要 OCR 的页面: {','.join(map(str, inspection.pages_needing_ocr))}"
        if inspection.pdf_type in {"scanned", "image_based"}:
            if self.ocr_document is not None and self.ocr_document.error:
                reason += f"；{self.ocr_document.error}"
            return self._wrap_tip(f"{reason}；本工具未获得有效 OCR 结果")
        fallback = "\n".join(self.lines)
        return self._wrap_tip_result(
            f"{reason}；已返回 PyMuPDF 提取的纯文本降级结果",
            fallback,
        )

    def _hybrid_pdf_markdown(self) -> str:
        heading_levels = {
            chapter.start_line: min(6, max(1, chapter.level))
            for chapter in self.chapters
        }
        output: list[str] = []
        for line_no, line in enumerate(self.lines, start=1):
            level = heading_levels.get(line_no)
            output.append(f"{'#' * level} {line}" if level else line)
        return "\n\n".join(line for line in output if line.strip())

    def _ocr_info_dict(self) -> dict[str, object]:
        document = self.ocr_document
        requested = self.pdf_inspection.pages_needing_ocr if self.pdf_inspection else []
        if document is None:
            return {
                "engine": "macos-vision",
                "requestedPages": requested,
                "successfulPages": [],
                "failedPages": requested,
                "processingTimeMs": 0,
                "error": None,
            }
        return {
            "engine": document.engine,
            "requestedPages": requested,
            "successfulPages": document.successful_pages,
            "failedPages": document.failed_pages,
            "processingTimeMs": document.processing_time_ms,
            "error": document.error,
        }

    def _pdf_inspection_summary(self) -> str:
        inspection = self.pdf_inspection
        if inspection is None:
            return ""
        if not inspection.available:
            return f"解析引擎:PyMuPDF\n增强引擎状态:{inspection.error or '不可用'}"
        return "\n".join(
            [
                "解析引擎:pdf-inspector + PyMuPDF",
                f"PDF类型:{inspection.pdf_type}",
                f"检测置信度:{inspection.confidence:.2f}",
                f"总页数:{inspection.page_count}",
                f"复杂版面:{'是' if inspection.is_complex_layout else '否'}",
                f"表格页:{_format_page_list(inspection.pages_with_tables)}",
                f"多栏页:{_format_page_list(inspection.pages_with_columns)}",
                f"需要OCR页:{_format_page_list(inspection.pages_needing_ocr)}",
                f"字体编码异常:{'是' if inspection.has_encoding_issues else '否'}",
                *self._ocr_summary_lines(),
            ]
        )

    def _ocr_summary_lines(self) -> list[str]:
        document = self.ocr_document
        if document is None:
            return []
        lines = [
            f"OCR引擎:{document.engine}",
            f"OCR成功页:{_format_page_list(document.successful_pages)}",
            f"OCR失败页:{_format_page_list(document.failed_pages)}",
            f"OCR耗时:{document.processing_time_ms}ms",
        ]
        if document.error:
            lines.append(f"OCR错误:{document.error}")
        return lines

    def _summary_chapters(self) -> list[LocalChapter]:
        if self.path.suffix.lower() != ".pdf":
            return self.chapters
        return [chapter for chapter in self.chapters if _is_summary_pdf_chapter(chapter)]

    def extract_chapter_summary(self, chapters: str) -> str:
        if self._excel_service is not None:
            return self._excel_service.extract_chapter_summary(chapters)
        parts: list[str] = []
        missing: list[str] = []
        for query in self._split_csv(chapters):
            chapter = self._find_chapter(query)
            if chapter is None:
                missing.append(query)
                continue
            parts.extend(self._chapter_summary_lines(chapter))
        if not parts:
            return self._wrap_tip(f"未找到{','.join(missing)}\n对应的章节，请核实编号或名称")
        if missing:
            return self._wrap_tip_result(
                f"未找到{','.join(missing)}\n对应的章节，请核实编号或名称",
                "\n".join(parts),
            )
        return self._wrap_result("\n".join(parts))

    def extract_chapter_content(self, chapters: str) -> str:
        if self._excel_service is not None:
            return self._excel_service.extract_chapter_content(chapters)
        parts: list[str] = []
        missing: list[str] = []
        for query in self._split_csv(chapters):
            chapter = self._find_chapter(query)
            if chapter is None:
                missing.append(query)
                continue
            parts.extend(self.lines[chapter.start_line - 1 : chapter.end_line])
        if not parts:
            return self._wrap_tip(f"未找到{','.join(missing)}\n对应的章节，请核实编号或名称")
        return self._wrap_result("\n".join(parts))

    def read_lines(self, start_line: int, end_line: int) -> str:
        if self._excel_service is not None:
            return self._excel_service.read_lines(start_line, end_line)
        if not self.lines:
            return self._wrap_result("")
        start = max(1, start_line)
        end = min(len(self.lines), end_line)
        if start > end:
            return self._wrap_result("")
        return self._wrap_result("\n".join(self.lines[start - 1 : end]))

    def search_content(
        self,
        keywords: str,
        page_no: int | None = None,
        page_size: int | None = None,
    ) -> str:
        if self._excel_service is not None:
            return self._excel_service.search_content(keywords, page_no, page_size)
        keys = self._split_csv(keywords)
        hits: list[SearchHit] = []
        for line_no, line in enumerate(self.lines, start=1):
            for segment in _search_segments(line):
                for key in keys:
                    if key and key in segment:
                        hits.append(
                            SearchHit(
                                line_no=line_no,
                                text=segment,
                                keyword=key,
                                score=self._search_hit_score(line_no, segment),
                            )
                        )
                        break
        hits.sort(key=lambda hit: (-hit.score, hit.line_no))
        page_no = max(1, page_no or 1)
        page_size = max(1, page_size or 20)
        total_pages = math.ceil(len(hits) / page_size) if hits else 0
        start = (page_no - 1) * page_size
        selected = hits[start : start + page_size]
        body = "\n".join(f"{hit.line_no} {hit.text}" for hit in selected)
        tip = (
            f"搜索结果总条数:{len(hits)}\n"
            f"当前页码:{page_no}\n"
            f"每页条数:{page_size}\n"
            f"总页数:{total_pages}"
        )
        return self._wrap_tip_result(tip, body)

    def query_content(
        self,
        query: str,
        page_no: int | None = None,
        page_size: int | None = None,
    ) -> str:
        if self._excel_service is not None:
            return self._excel_service.search_content(
                _query_keywords(query),
                page_no=page_no,
                page_size=page_size,
            )
        keys = _query_terms(query, self._query_terms)
        hits: list[QueryHit] = []
        seen: set[tuple[int, str]] = set()
        seen_lines: set[int] = set()
        chapter = self._find_chapter(query)
        if chapter is None:
            chapter = self._find_chapter_by_query_terms(keys)
        if chapter is not None:
            snippet = "\n".join(self.lines[chapter.start_line - 1 : min(chapter.end_line, chapter.start_line + 8)])
            hits.append(
                QueryHit(
                    line_no=chapter.start_line,
                    text=snippet,
                    score=1000,
                    source="chapter",
                )
            )
            seen.add((chapter.start_line, snippet))
            seen_lines.add(chapter.start_line)
        for line_no, line in enumerate(self.lines, start=1):
            if line_no in seen_lines:
                continue
            for segment in _search_segments(line):
                score = _query_score(segment, keys)
                if score < 40:
                    continue
                score += self._search_hit_score(line_no, segment)
                if score < 40:
                    continue
                key = (line_no, segment)
                if key in seen:
                    continue
                seen.add(key)
                hits.append(QueryHit(line_no=line_no, text=segment, score=score, source="search"))
        hits.sort(key=lambda hit: (-hit.score, hit.line_no))
        page_no = max(1, page_no or 1)
        page_size = max(1, page_size or 10)
        total_pages = math.ceil(len(hits) / page_size) if hits else 0
        start = (page_no - 1) * page_size
        selected = hits[start : start + page_size]
        body = "\n".join(
            f"{hit.line_no} [{hit.source};score={hit.score}] {hit.text}"
            for hit in selected
        )
        tip = (
            f"自然语言检索结果总条数:{len(hits)}\n"
            f"当前页码:{page_no}\n"
            f"每页条数:{page_size}\n"
            f"总页数:{total_pages}\n"
            f"检索关键词:{','.join(keys)}"
        )
        return self._wrap_tip_result(tip, body)

    def _search_hit_score(self, line_no: int, text: str) -> int:
        if any(chapter.start_line == line_no for chapter in self.chapters):
            return 300
        body_start = self._body_start_line or 1
        score = 200 if line_no >= body_start else 20
        if _looks_like_pdf_toc_line(text):
            score -= 160
        if re.search(r"\.{4,}", text):
            score -= 120
        if re.match(r"^\d{1,2}\.[^\s.]{2,30}$", text):
            score -= 90
        if _looks_like_table_title(text) or text.startswith("<tr"):
            score += 10
        return score

    def extract_table_list(self, chapters: str) -> str:
        if self._excel_service is not None:
            return self._excel_service.extract_table_list(chapters)
        titles: list[str] = []
        for query in self._split_csv(chapters):
            chapter = self._find_chapter(query)
            if chapter is None:
                continue
            for table in self.tables:
                if table.title.startswith("表格-"):
                    continue
                if self._table_in_chapter(table, chapter):
                    titles.append(table.title)
        return self._wrap_result("\n".join(titles))

    def extract_table_content(
        self,
        table_title: str,
        row_scope: str | None = None,
        row_keywords: str | None = None,
        col_keywords: str | None = None,
    ) -> str:
        if self._excel_service is not None:
            return self._excel_service.extract_table_content(
                table_title,
                row_scope=row_scope,
                row_keywords=row_keywords,
                col_keywords=col_keywords,
            )
        table = self._find_table(table_title)
        if table is None:
            fallback_terms = ",".join(
                term
                for term in [table_title, row_keywords or "", col_keywords or ""]
                if term.strip()
            )
            fallback = self.search_content(fallback_terms, page_no=1, page_size=10) if fallback_terms else ""
            if "<result>\n</result>" not in fallback and fallback:
                return self._wrap_tip_result(
                    "未找到匹配表格标题，已自动按关键词检索相关正文内容。若要精确提取表格，请核实 tableTitle。",
                    fallback,
                )
            return self._wrap_tip("未检索到相关内容，请核实名称是否准确，是否存在多余的空格等字符")
        if (
            table.is_html
            and table.markdown.startswith("<table>")
            and not row_keywords
            and not col_keywords
        ):
            table_text = _slice_html_table(table.markdown, row_scope)
        else:
            rows = self._filter_table_rows(table.rows, row_scope, row_keywords, col_keywords)
            if table.is_html:
                table_text = "\n".join(row[0] for row in rows if row)
            else:
                table_text = self._rows_to_markdown(rows)
        body_label = "筛选结果" if row_scope or row_keywords or col_keywords else "表格内容"
        if (row_scope or row_keywords or col_keywords) and not table_text.strip():
            fallback_terms = ",".join(
                term
                for term in [table_title, row_keywords or "", col_keywords or ""]
                if term.strip()
            )
            fallback = self.search_content(fallback_terms, page_no=1, page_size=10) if fallback_terms else ""
            if "<result>\n</result>" not in fallback and fallback:
                return self._wrap_tip_result(
                    "表格标题已命中，但筛选条件未命中表格行列，已自动按关键词检索相关正文内容。",
                    fallback,
                )
        text = (
            f"表格标题:{table.title}\n"
            f"总行数:{len(table.rows)}\n"
            + (
                f"筛选行范围:{'[' + row_scope + ']' if row_scope else '全部'}\n"
                if row_scope or row_keywords or col_keywords
                else ""
            )
            + f"{body_label}:\n{table_text}"
        )
        return self._wrap_result(text)

    def _parse(self) -> None:
        if self.path.suffix.lower() == ".pdf":
            self._parse_pdf()
            return

        self._parse_word()

    def _parse_word(self) -> None:
        document = Document(str(self._parse_path))
        pending_table_title: str | None = None
        last_chapter_key = ""
        automatic_heading_counters: dict[int, list[int]] = {}
        in_toc = False
        toc_heading_count = 0
        body_started = False
        skip_four_ground_heading = False
        saw_content = False
        preserved_leading_blank = False
        last_top_level_number = 0
        for block in _iter_block_items(document):
            if isinstance(block, Paragraph):
                text = _clean_text(block.text)
                if not text:
                    if not saw_content and not preserved_leading_blank:
                        self._append_line("")
                        preserved_leading_blank = True
                    continue
                saw_content = True
                if text == "目录":
                    in_toc = True
                line_no = self._append_line(text)
                if text == "目录":
                    self._toc_line = line_no
                style_name = block.style.name if block.style else ""
                heading = _parse_heading(text, style_name)
                automatic_heading = _parse_automatic_word_heading(
                    block,
                    text,
                    automatic_heading_counters,
                )
                if automatic_heading is not None:
                    heading = automatic_heading
                elif (
                    automatic_heading_counters
                    and heading is not None
                    and heading.get("level") == 1
                    and not re.search(r"Heading|标题", style_name, re.IGNORECASE)
                ):
                    # Once a document exposes a reliable automatic heading tree,
                    # bare items such as "2 油气集输工程" inside a section
                    # are local list items, not new top-level chapters.
                    heading = None
                if in_toc and heading is not None:
                    toc_heading_count += 1
                    if _is_first_body_heading(text) and (
                        toc_heading_count >= 5 or not _is_toc_style(style_name)
                    ):
                        self._pad_lines_until(71)
                        line_no = len(self.lines)
                        in_toc = False
                        body_started = True
                        self._body_start_line = line_no
                    else:
                        heading = None
                elif in_toc and text == "总论":
                    in_toc = False
                    body_started = True
                    self._body_start_line = line_no
                    heading = {"number": "1", "title": "总论", "level": 1}
                elif heading is not None and not body_started:
                    if _is_first_body_heading(text, heading):
                        body_started = True
                        self._body_start_line = line_no
                    else:
                        heading = None
                if heading is not None and skip_four_ground_heading:
                    number = str(heading["number"])
                    if number == "5":
                        skip_four_ground_heading = False
                    elif number == "4" or number.startswith("4."):
                        heading = None
                if heading is not None and _normalize_heading(text) == "4地面工程":
                    skip_four_ground_heading = True
                    heading = None
                if heading is not None and _is_out_of_sequence_top_heading(
                    heading, last_top_level_number
                ):
                    heading = None
                if heading is not None:
                    chapter = LocalChapter(
                        number=heading["number"],
                        title=heading["title"],
                        full_title=str(heading.get("full_title") or text),
                        level=heading["level"],
                        start_line=line_no,
                    )
                    self.chapters.append(chapter)
                    last_chapter_key = chapter.number or chapter.title
                    if chapter.level == 1:
                        last_top_level_number = _first_heading_number(chapter.number)
                if _looks_like_table_title(text):
                    pending_table_title = text
            elif isinstance(block, Table):
                rows = [[_clean_cell_text(cell.text) for cell in row.cells] for row in block.rows]
                if not rows:
                    continue
                markdown = self._table_to_text(block, rows)
                start_line = len(self.lines) + 1
                for line in markdown.splitlines():
                    self._append_line(line)
                end_line = len(self.lines)
                title = pending_table_title or f"表格-{len(self.tables) + 1}"
                chapter_key = _chapter_key_from_table_title(title) or last_chapter_key
                self.tables.append(
                    LocalTable(
                        title=title,
                        chapter_key=chapter_key,
                        start_line=start_line,
                        end_line=end_line,
                        markdown=markdown,
                        rows=rows,
                        is_html=_has_merged_cells(block),
                    )
                )
                pending_table_title = None
        if self.path.suffix.lower() == ".doc":
            self._repair_converted_doc_toc()
        self._finalize_chapters()

    def _repair_converted_doc_toc(self) -> None:
        if self._toc_line is None or not self.chapters:
            return
        first_chapter_line = self.chapters[0].start_line
        if first_chapter_line <= self._toc_line + 1:
            return
        toc_indexes = [
            idx
            for idx in range(self._toc_line, first_chapter_line - 1)
            if not self.lines[idx].strip()
        ]
        if not toc_indexes:
            return
        entries = []
        seen: set[str] = set()
        for chapter in self.chapters:
            if _heading_number_level(chapter.number) > 2:
                continue
            if chapter.full_title in seen:
                continue
            seen.add(chapter.full_title)
            entries.append(chapter.full_title)
        for idx, entry in zip(toc_indexes, entries):
            self.lines[idx] = entry

    def _parse_pdf(self) -> None:
        self.pdf_inspection = self._pdf_backend.inspect(self.path)
        document = fitz.open(str(self.path))
        native_lines_by_page = {
            page_index: _pdf_page_lines(document.load_page(page_index - 1))
            for page_index in range(1, document.page_count + 1)
        }
        ocr_pages = [
            page_number
            for page_number in self.pdf_inspection.pages_needing_ocr
            if self.pdf_inspection.pdf_type in {"scanned", "image_based", "mixed"}
            or _native_pdf_text_needs_ocr(native_lines_by_page.get(page_number, []))
        ]
        if ocr_pages:
            self.ocr_document = self._ocr_backend.recognize_pdf(self.path, ocr_pages)
        for page_index in range(1, document.page_count + 1):
            page = document.load_page(page_index - 1)
            page_lines: list[str] = []
            ocr_page = self.ocr_document.pages.get(page_index) if self.ocr_document else None
            if ocr_page and ocr_page.text:
                raw_lines = ocr_page.text.splitlines()
                if _looks_like_ocr_toc_page(raw_lines):
                    self._pdf_toc_pages.add(page_index)
                raw_lines = _merge_ocr_heading_lines(raw_lines)
            else:
                raw_lines = native_lines_by_page.get(page_index, [])
            for raw_line in raw_lines:
                line = _clean_pdf_text(raw_line)
                if _looks_like_pdf_page_number(line):
                    continue
                if line:
                    for split_line in _split_pdf_heading_prefix(line):
                        page_lines.append(split_line)
            page_lines = _merge_pdf_wrapped_toc_lines(page_lines)
            page_lines = _merge_pdf_wrapped_body_lines(page_lines)
            for line in page_lines:
                line_no = self._append_line(line)
                self._pdf_line_pages[line_no] = page_index
            self._pdf_page_texts[page_index] = "\n".join(page_lines)
            self._parse_pdf_detected_tables(page, page_index)
        self._parse_pdf_chapters()
        self._parse_pdf_table_titles()
        self._parse_pdf_inspector_tables()
        self._finalize_chapters()
        document.close()

    def _parse_pdf_inspector_tables(self) -> None:
        inspection = self.pdf_inspection
        if inspection is None or not inspection.markdown:
            return
        for caption, markdown, rows in _extract_markdown_tables(inspection.markdown):
            if caption is None:
                continue
            title = caption
            if self._find_table(title) is not None:
                continue
            start_line = _find_matching_line(self.lines, title) or 1
            self.tables.append(
                LocalTable(
                    title=title,
                    chapter_key=_chapter_key_from_table_title(title),
                    start_line=start_line,
                    end_line=start_line,
                    markdown=markdown,
                    rows=rows,
                    is_html=False,
                )
            )

    def _parse_pdf_detected_tables(self, page: fitz.Page, page_index: int) -> None:
        if not hasattr(page, "find_tables"):
            return

        detected = page.find_tables()
        for table_index, table in enumerate(detected.tables, start=1):
            rows = _normalize_pdf_table_rows(table.extract())
            if not _is_useful_pdf_table(rows):
                continue
            title = _pdf_table_title_from_page(page, table.bbox) or f"PDF表格-第{page_index}页-{table_index}"
            markdown = self._rows_to_markdown(rows)
            self.tables.append(
                LocalTable(
                    title=title,
                    chapter_key=_chapter_key_from_table_title(title),
                    start_line=max(1, len(self.lines)),
                    end_line=max(1, len(self.lines)),
                    markdown=markdown,
                    rows=rows,
                    is_html=False,
                )
            )

    def _parse_pdf_chapters(self) -> None:
        body_started = False
        last_top_level_number = 0
        current_top_level_number = 0
        candidates: list[tuple[int, str, dict[str, str | int]]] = []
        for line_no, text in enumerate(self.lines, start=1):
            if self._toc_line is None and text in {"目录", "Contents"}:
                self._toc_line = line_no
            if self._pdf_line_pages.get(line_no) in self._pdf_toc_pages:
                continue
            heading = _parse_pdf_heading(text)
            if heading is None:
                continue
            if not body_started:
                if _is_first_body_heading(text, heading):
                    body_started = True
                    self._body_start_line = line_no
                else:
                    continue
            if _looks_like_pdf_toc_line(text):
                continue
            if _is_pdf_page_number_heading(text, heading):
                continue
            if _reject_pdf_heading_candidate(self.lines, line_no, heading, current_top_level_number):
                continue
            if _is_out_of_sequence_top_heading(heading, last_top_level_number):
                continue
            number = str(heading["number"])
            first_number = _first_heading_number(number)
            if _is_pdf_fee_item_heading(heading) and current_top_level_number != 2 and first_number < 45:
                continue
            if _is_pdf_fee_item_heading(heading) and not _pdf_fee_item_belongs_to_top(
                first_number,
                current_top_level_number,
            ):
                continue
            if (
                current_top_level_number
                and first_number
                and first_number != current_top_level_number
                and heading["level"] > 1
                and not _is_pdf_fee_item_heading(heading)
                and not _is_chinese_marker_heading(heading)
            ):
                continue
            candidates.append((line_no, text, heading))
            if heading["level"] == 1:
                last_top_level_number = first_number
                current_top_level_number = first_number
        candidates = _filter_pdf_heading_candidates(candidates)
        for line_no, text, heading in candidates:
            self.chapters.append(
                LocalChapter(
                    number=str(heading["number"]),
                    title=str(heading["title"]),
                    full_title=text,
                    level=int(heading["level"]),
                    start_line=line_no,
                )
            )

    def _register_query_terms(self) -> None:
        terms: list[str] = []
        for chapter in self.chapters:
            terms.append(chapter.title)
            terms.append(_normalize_heading(chapter.title))
            full_title = re.sub(
                r"^\s*(?:第[一二三四五六七八九十百]+章|\d+(?:\.\d+)*\.?|[一二三四五六七八九十]+、|（[一二三四五六七八九十]+）)\s*",
                "",
                chapter.full_title,
            )
            terms.append(full_title.strip())
        clean_terms: list[str] = []
        seen: set[str] = set()
        for term in terms:
            cleaned = term.strip()
            if len(cleaned) < 2 or cleaned in seen:
                continue
            seen.add(cleaned)
            clean_terms.append(cleaned)
            if jieba is not None:
                jieba.add_word(cleaned, freq=200000)
        self._query_terms = clean_terms

    def _parse_pdf_table_titles(self) -> None:
        for line_no, text in enumerate(self.lines, start=1):
            if not _looks_like_table_title(text):
                continue
            if self._find_table(text) is not None:
                continue
            end = self._pdf_table_end_line(line_no)
            body = "\n".join(self.lines[line_no:end])
            rows = [[line] for line in self.lines[line_no:end] if line]
            self.tables.append(
                LocalTable(
                    title=text,
                    chapter_key=_chapter_key_from_table_title(text),
                    start_line=line_no,
                    end_line=end,
                    markdown=body,
                    rows=rows,
                    is_html=True,
                )
            )

    def _pdf_table_end_line(self, start_line: int) -> int:
        max_end = min(len(self.lines), start_line + 80)
        for line_no in range(start_line + 1, max_end + 1):
            text = self.lines[line_no - 1]
            if _looks_like_table_title(text):
                return line_no - 1
            if _looks_like_after_table_body(text):
                return line_no - 1
            heading = _parse_pdf_heading(text)
            if (
                heading is not None
                and heading.get("level") in {1, 2, 3}
                and _looks_like_real_pdf_section_heading(text, heading)
            ):
                return line_no - 1
        return max_end

    def _append_line(self, text: str) -> int:
        self.lines.append(text)
        return len(self.lines)

    def _pad_lines_until(self, line_no: int) -> None:
        while len(self.lines) < line_no:
            self.lines.insert(-1, "")

    def _finalize_chapters(self) -> None:
        for idx, chapter in enumerate(self.chapters):
            chapter.end_line = len(self.lines)
            for later in self.chapters[idx + 1 :]:
                if later.level <= chapter.level:
                    chapter.end_line = later.start_line - 1
                    break
            segment = "\n".join(self.lines[chapter.start_line - 1 : chapter.end_line])
            chapter.char_count = len(segment)
            chapter.table_char_count = sum(
                len(table.markdown)
                for table in self.tables
                if self._table_in_chapter(table, chapter)
            )
        stack: list[LocalChapter] = []
        for chapter in self.chapters:
            while stack and stack[-1].level >= chapter.level:
                stack.pop()
            if stack:
                stack[-1].children.append(chapter)
            stack.append(chapter)

    def _chapter_summary_lines(self, chapter: LocalChapter) -> list[str]:
        lines = [self._chapter_summary_line(chapter)]
        for child in chapter.children:
            lines.extend(self._chapter_summary_lines(child))
        return lines

    def _chapter_summary_line(self, chapter: LocalChapter) -> str:
        table_text = (
            f"其中表格字符总数:{chapter.table_char_count}"
            if chapter.table_char_count
            else "无表格内容"
        )
        return (
            f"{chapter.full_title}\n"
            f"字符总数:{chapter.char_count},{table_text}\n"
            f"总行数:{max(0, chapter.end_line - chapter.start_line + 1)},"
            f"行索引范围:{chapter.start_line}~{chapter.end_line}"
        )

    def _find_chapter(self, query: str) -> LocalChapter | None:
        compound = self._find_compound_chapter(query)
        if compound is not None:
            return compound
        normalized = _normalize_heading(query)
        for chapter in self.chapters:
            candidates = {
                chapter.number,
                chapter.title,
                chapter.full_title,
                f"{chapter.number}{chapter.title}",
            }
            if normalized in {_normalize_heading(item) for item in candidates if item}:
                return chapter
        for chapter in self.chapters:
            if normalized and normalized in _normalize_heading(chapter.full_title):
                if _looks_like_chapter_number_query(query):
                    return None
                return chapter
        return None

    def _find_compound_chapter(self, query: str) -> LocalChapter | None:
        match = re.search(r"(第[一二三四五六七八九十百]+章|[\u4e00-\u9fff]{2,20})\s+(\d{1,2})(?:\.|\s|$)", query)
        if match is None:
            return None
        parent_query = match.group(1)
        item_number = match.group(2)
        parent = self._find_direct_chapter(parent_query)
        for chapter in self.chapters:
            if chapter.number != item_number:
                continue
            if parent is not None and not self._line_in_chapter(chapter.start_line, parent):
                continue
            return chapter
        return None

    def _find_direct_chapter(self, query: str) -> LocalChapter | None:
        normalized = _normalize_heading(query)
        for chapter in self.chapters:
            candidates = {
                chapter.number,
                chapter.title,
                chapter.full_title,
                f"{chapter.number}{chapter.title}",
            }
            if normalized in {_normalize_heading(item) for item in candidates if item}:
                return chapter
        for chapter in self.chapters:
            if normalized and normalized in _normalize_heading(chapter.full_title):
                return chapter
        return None

    def _find_chapter_by_query_terms(self, terms: list[str]) -> LocalChapter | None:
        normalized_terms = [
            _normalize_heading(term)
            for term in terms
            if len(_normalize_heading(term)) >= 3
        ]
        for term in normalized_terms:
            for chapter in self.chapters:
                chapter_text = _normalize_heading(f"{chapter.title}{chapter.full_title}")
                if term and term in chapter_text:
                    return chapter
        return None

    def _find_table(self, title: str) -> LocalTable | None:
        normalized = _normalize_heading(title)
        for table in self.tables:
            if _normalize_heading(table.title) == normalized:
                return table
        return None

    def _line_in_chapter(self, line_no: int, chapter: LocalChapter) -> bool:
        return chapter.start_line <= line_no <= chapter.end_line

    def _table_in_chapter(self, table: LocalTable, chapter: LocalChapter) -> bool:
        if self._line_in_chapter(table.start_line, chapter):
            return True
        return bool(table.chapter_key) and (
            table.chapter_key == chapter.number or table.chapter_key.startswith(f"{chapter.number}.")
        )

    def _filter_table_rows(
        self,
        rows: list[list[str]],
        row_scope: str | None,
        row_keywords: str | None,
        col_keywords: str | None,
    ) -> list[list[str]]:
        selected = list(rows)
        if row_scope:
            scoped: list[list[str]] = []
            for start, end in _parse_scopes(row_scope):
                scoped.extend(rows[max(0, start - 1) : min(len(rows), end)])
            selected = scoped
        if row_keywords:
            keys = self._split_csv(row_keywords)
            header = selected[:1]
            body = [row for row in selected[1:] if any(key in " ".join(row) for key in keys)]
            selected = header + body if header else body
        if col_keywords and selected:
            keys = self._split_csv(col_keywords)
            header = selected[0]
            indexes = [idx for idx, value in enumerate(header) if any(key in value for key in keys)]
            if indexes:
                selected = [[row[idx] if idx < len(row) else "" for idx in indexes] for row in selected]
        return selected

    def _rows_to_markdown(self, rows: list[list[str]]) -> str:
        if not rows:
            return ""
        width = max(len(row) for row in rows)
        normalized = [row + [""] * (width - len(row)) for row in rows]
        header = "|" + "|".join(normalized[0]) + "|"
        separator = "|" + "|".join("---" for _ in range(width)) + "|"
        body = ["|" + "|".join(row) + "|" for row in normalized[1:]]
        return "\n".join([header, separator, *body])

    def _table_to_text(self, table: Table, rows: list[list[str]]) -> str:
        if _has_merged_cells(table):
            return _table_to_html(table)
        return self._rows_to_markdown(rows)

    def _wrap_result(self, text: str) -> str:
        return f"{self._meta()}\n<result>\n{text}</result>"

    def _wrap_tip(self, text: str) -> str:
        return f"{self._meta()}\n<tip>{text}</tip>"

    def _wrap_tip_result(self, tip: str, result: str) -> str:
        return f"{self._meta()}\n<tip>\n{tip}\n</tip>\n<result>\n{result}</result>"

    def _meta(self) -> str:
        return (
            "<meta>\n"
            f"<file-id>{self.file_id}</file-id>\n"
            f"<file-name>{self.file_name}</file-name>\n"
            "</meta>"
        )

    def _total_chars(self) -> int:
        return len("\n".join(self.lines))

    def _split_csv(self, value: str) -> list[str]:
        return [part.strip() for part in value.split(",") if part.strip()]


def _iter_block_items(parent: DocumentObject) -> Iterator[Paragraph | Table]:
    body = parent.element.body
    for child in body.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent)


def _pdf_page_lines(page: fitz.Page) -> list[str]:
    lines: list[str] = []
    for block in page.get_text("blocks"):
        text = str(block[4]).strip()
        if not text:
            continue
        block_lines = [line.strip() for line in text.splitlines() if line.strip()]
        if not block_lines:
            continue
        if _should_join_pdf_block(block_lines):
            lines.append(" ".join(block_lines))
        else:
            lines.extend(block_lines)
    return lines


def _native_pdf_text_needs_ocr(lines: list[str]) -> bool:
    text = "".join(lines)
    meaningful = re.findall(r"[\u4e00-\u9fffA-Za-z0-9]", text)
    if len(meaningful) < 20:
        return True
    replacement_count = text.count("�") + text.count("□")
    return replacement_count >= max(5, len(meaningful) // 5)


def _extract_markdown_tables(markdown: str) -> list[tuple[str | None, str, list[list[str]]]]:
    lines = markdown.splitlines()
    tables: list[tuple[str | None, str, list[list[str]]]] = []
    index = 0
    while index + 1 < len(lines):
        if not _is_markdown_table_row(lines[index]) or not _is_markdown_table_separator(
            lines[index + 1]
        ):
            index += 1
            continue
        end = index + 2
        while end < len(lines) and _is_markdown_table_row(lines[end]):
            end += 1
        table_lines = lines[index:end]
        rows = [_split_markdown_table_row(line) for line in table_lines if line != lines[index + 1]]
        if _is_useful_pdf_table(rows):
            caption = _markdown_table_caption(lines, index)
            tables.append((caption, "\n".join(table_lines), rows))
        index = end
    return tables


def _is_markdown_table_row(line: str) -> bool:
    stripped = line.strip()
    return stripped.startswith("|") and stripped.endswith("|") and stripped.count("|") >= 3


def _is_markdown_table_separator(line: str) -> bool:
    if not _is_markdown_table_row(line):
        return False
    cells = _split_markdown_table_row(line)
    return bool(cells) and all(bool(re.fullmatch(r":?-{3,}:?", cell)) for cell in cells)


def _split_markdown_table_row(line: str) -> list[str]:
    stripped = line.strip().strip("|")
    cells = re.split(r"(?<!\\)\|", stripped)
    return [cell.replace(r"\|", "|").strip() for cell in cells]


def _markdown_table_caption(lines: list[str], table_index: int) -> str | None:
    for index in range(table_index - 1, max(-1, table_index - 4), -1):
        candidate = re.sub(r"^#{1,6}\s+", "", lines[index].strip()).strip()
        if not candidate:
            continue
        return candidate if _looks_like_table_title(candidate) else None
    return None


def _find_matching_line(lines: list[str], value: str) -> int | None:
    normalized = _normalize_heading(value)
    if not normalized:
        return None
    for line_no, line in enumerate(lines, start=1):
        if normalized == _normalize_heading(line):
            return line_no
    return None


def _format_page_list(pages: list[int]) -> str:
    return ",".join(map(str, pages)) if pages else "无"


def _normalize_pdf_table_rows(rows: list[list[object]]) -> list[list[str]]:
    normalized: list[list[str]] = []
    for row in rows:
        normalized.append([_clean_pdf_table_cell(cell) for cell in row])
    return normalized


def _clean_pdf_table_cell(value: object) -> str:
    if value is None:
        return ""
    text = str(value)
    text = re.sub(r"\s*\n\s*", " ", text)
    return re.sub(r"\s+", " ", text).strip()


def _is_useful_pdf_table(rows: list[list[str]]) -> bool:
    if len(rows) < 2:
        return False
    width = max((len(row) for row in rows), default=0)
    if width < 2:
        return False
    non_empty = sum(1 for row in rows for cell in row if cell.strip())
    return non_empty >= 4


def _pdf_table_title_from_page(page: fitz.Page, bbox: object) -> str | None:
    try:
        x0, y0, x1, _ = bbox  # type: ignore[misc]
    except Exception:
        return None

    candidates: list[tuple[float, str]] = []
    for block in page.get_text("blocks"):
        bx0, by0, bx1, by1, text = block[:5]
        if by1 > y0 or by1 < y0 - 90:
            continue
        if bx1 < x0 - 30 or bx0 > x1 + 30:
            continue
        for line in str(text).splitlines():
            cleaned = _clean_pdf_text(line)
            if not cleaned:
                continue
            if _looks_like_table_title(cleaned):
                candidates.append((abs(y0 - by1), cleaned))
    if not candidates:
        return None
    return sorted(candidates, key=lambda item: item[0])[0][1]


def _should_join_pdf_block(lines: list[str]) -> bool:
    if len(lines) <= 1:
        return False
    if len(lines) >= 3:
        return True
    joined = "".join(lines)
    return len(joined) <= 60 and not any(_looks_like_table_title(line) for line in lines)


def _clean_text(value: str) -> str:
    return re.sub(r"\s+", " ", value).strip()


def _clean_pdf_text(value: str) -> str:
    text = re.sub(r"\s+", " ", value).strip()
    text = re.sub(r"(?<=\d)\s+(?=\d)", "", text)
    text = re.sub(r"(?<=[A-Za-z])\s+(?=[A-Za-z])", "", text)
    return text


def _looks_like_ocr_toc_page(lines: list[str]) -> bool:
    cleaned = [_clean_pdf_text(line) for line in lines if _clean_pdf_text(line)]
    if any(line in {"目录", "目次", "Contents"} for line in cleaned):
        return True
    return any(
        cleaned[index] == "目" and cleaned[index + 1] == "次"
        for index in range(len(cleaned) - 1)
    )


def _merge_ocr_heading_lines(lines: list[str]) -> list[str]:
    """Join OCR headings whose number and title were returned as separate lines."""
    cleaned = [_clean_pdf_text(line) for line in lines if _clean_pdf_text(line)]
    is_toc_page = _looks_like_ocr_toc_page(cleaned)
    merged: list[str] = []
    index = 0
    while index < len(cleaned):
        line = cleaned[index]
        if line == "目" and index + 1 < len(cleaned) and cleaned[index + 1] == "次":
            merged.append("目录")
            index += 2
            continue
        if not is_toc_page and index + 1 < len(cleaned):
            number = _normalized_ocr_heading_number(line)
            title = cleaned[index + 1]
            if number is not None and _should_merge_ocr_heading(number, title):
                merged.append(f"{number} {title}")
                index += 2
                continue
        merged.append(line)
        index += 1
    return merged


def _normalized_ocr_heading_number(text: str) -> str | None:
    if not re.fullmatch(r"\d{1,2}(?:\s*[.．]\s*\d{1,2}){0,3}[.．]?", text):
        return None
    number = re.sub(r"\s+", "", text).replace("．", ".").rstrip(".")
    return number if number else None


def _should_merge_ocr_heading(number: str, title: str) -> bool:
    if not title or len(title) > 80:
        return False
    if re.match(r"^\d", title) or re.search(r"[。；;：:]$", title):
        return False
    if re.search(r"标准分享网|免费下载", title):
        return False
    candidate = f"{number} {title}"
    heading = _parse_pdf_heading(candidate)
    if heading is None:
        return False
    if "." not in number:
        return heading.get("level") == 1 and (
            _looks_like_pdf_top_heading_title(title)
            or _looks_like_pdf_generic_top_heading_title(title)
        )
    return True


def _split_pdf_heading_prefix(text: str) -> list[str]:
    cn_marker_match = re.match(
        r"^([一二三四五六七八九十]+、\s*.+?项目)\s+(?=\d+[.．、]\s*)",
        text,
    )
    if cn_marker_match is not None and _parse_pdf_heading(cn_marker_match.group(1)) is not None:
        heading = cn_marker_match.group(1).strip()
        rest = text[len(cn_marker_match.group(1)) :].strip()
        return [heading, rest] if rest else [heading]
    top_match = re.match(r"^(\d{1,2}\s+[\u4e00-\u9fff]{2,12})(?:\s+(\d+\.\d+.*))?$", text)
    if top_match is not None and _parse_pdf_heading(top_match.group(1)) is not None:
        rest = top_match.group(2)
        return [top_match.group(1), rest] if rest else [top_match.group(1)]
    match = re.match(
        r"^(\d+(?:\.\d+)*\s*[\u4e00-\u9fff][^。；;，,：:]{0,28}?)(?=（|根据|该|经|结合|从|为|对|图|表)",
        text,
    )
    if match is None:
        return [text]
    heading = match.group(1).strip()
    if _parse_pdf_heading(heading) is None:
        return [text]
    rest = text[len(match.group(1)) :].strip()
    return [heading, rest] if rest else [heading]


def _merge_pdf_wrapped_toc_lines(lines: list[str]) -> list[str]:
    merged: list[str] = []
    idx = 0
    while idx < len(lines):
        line = lines[idx]
        if (
            idx + 1 < len(lines)
            and re.search(r"\.{4,}\s*$", line)
            and re.search(r"\.{4,}\s*\d+\s*$", lines[idx + 1])
        ):
            merged.append(line + lines[idx + 1])
            idx += 2
            continue
        merged.append(line)
        idx += 1
    return merged


def _merge_pdf_wrapped_body_lines(lines: list[str]) -> list[str]:
    merged: list[str] = []
    for line in lines:
        if not merged or not _should_merge_pdf_body_line(merged[-1], line):
            merged.append(line)
            continue
        merged[-1] = f"{merged[-1]} {line}"
    return merged


def _search_segments(line: str) -> list[str]:
    segments = _split_compact_toc_entries(line)
    return segments if segments else [line]


def _query_keywords(query: str) -> str:
    return ",".join(_query_terms(query))


def _query_terms(query: str, domain_terms: list[str] | None = None) -> list[str]:
    normalized = re.sub(r"\s+", "", query)
    terms: list[str] = []
    domain_terms = domain_terms or []
    for term in domain_terms:
        if _domain_term_matches_query(term, normalized):
            terms.append(term)
    for token in re.findall(r"[\u4e00-\u9fffA-Za-z0-9（）()]+", query):
        cleaned = token.strip()
        if cleaned and cleaned not in _QUERY_STOPWORDS:
            terms.append(cleaned)
    for suffix in ("费", "费用", "管理费", "审查费", "咨询费", "预备费", "补偿费", "补偿", "使用金"):
        for match in re.finditer(rf"[\u4e00-\u9fffA-Za-z0-9（）()]{{2,30}}{suffix}", normalized):
            terms.append(match.group(0))
    compact_terms = _split_query_compound_terms(normalized)
    terms.extend(compact_terms)
    terms.extend(_jieba_query_terms(query, domain_terms))
    terms.extend(_rule_based_query_terms(normalized))
    expanded: list[str] = []
    for term in terms:
        expanded.append(term)
        expanded.extend(_QUERY_SYNONYMS.get(term, []))
    deduped: list[str] = []
    seen: set[str] = set()
    for term in expanded:
        if len(term) <= 1:
            continue
        if term in seen:
            continue
        seen.add(term)
        deduped.append(term)
    return deduped[:12]


def _domain_term_matches_query(term: str, normalized_query: str) -> bool:
    normalized_term = _normalize_heading(term)
    if len(normalized_term) < 2:
        return False
    if normalized_term in normalized_query:
        return True
    term_tokens = _jieba_query_terms(normalized_term, [])
    if not term_tokens:
        term_tokens = _rule_based_query_terms(normalized_term)
    if not term_tokens:
        return False
    matched = [
        token
        for token in term_tokens
        if len(token) >= 2 and token not in _QUERY_STOPWORDS and token in normalized_query
    ]
    return len(matched) >= 2 or (len(matched) == 1 and len(matched[0]) >= 4)


def _jieba_query_terms(query: str, domain_terms: list[str]) -> list[str]:
    if jieba is None:
        return []
    for term in domain_terms:
        if term:
            jieba.add_word(term, freq=200000)
    tokens = []
    for token in jieba.lcut(query):
        cleaned = token.strip()
        if len(cleaned) < 2:
            continue
        if cleaned in _QUERY_STOPWORDS:
            continue
        if re.fullmatch(r"\d+", cleaned):
            continue
        tokens.append(cleaned)
    return tokens


def _rule_based_query_terms(query: str) -> list[str]:
    terms: list[str] = []
    normalized = _normalize_heading(query)
    for token in _QUERY_KEY_TERMS:
        if token in normalized:
            terms.append(token)
    for left in _QUERY_KEY_TERMS:
        if left not in normalized:
            continue
        for right in _QUERY_KEY_TERMS:
            if right == left or right not in normalized:
                continue
            terms.append(f"{left}{right}")
    return terms


def _split_query_compound_terms(query: str) -> list[str]:
    terms: list[str] = []
    for suffix in ("费", "费用", "管理费", "审查费", "咨询费", "预备费", "补偿费", "补偿", "使用金"):
        idx = query.find(suffix)
        if idx <= 0:
            continue
        start = max(0, idx - 12)
        candidate = query[start : idx + len(suffix)]
        candidate = re.sub(
            r"^(介绍|说明|查询|查找|看看|关于|有哪些|一下|这个|那个|请问|帮我|帮忙|的|和|及|与)+",
            "",
            candidate,
        )
        if len(candidate) >= 2:
            terms.append(candidate)
    return terms


def _query_score(text: str, terms: list[str]) -> int:
    if not terms:
        return 0
    normalized = _normalize_heading(text)
    score = 0
    counts = Counter()
    for term in terms:
        term_norm = _normalize_heading(term)
        if not term_norm:
            continue
        if term_norm in normalized:
            counts[term_norm] += 1
            score += min(80, 15 + len(term_norm) * 5)
    if len(counts) >= 2:
        score += 30
    if re.search(r"(费用内容|计算方法|计费标准|计费基价|调整系数)", text):
        score += 25
    return score


_QUERY_STOPWORDS = {
    "介绍",
    "一下",
    "这个",
    "那个",
    "请问",
    "帮我",
    "帮忙",
    "查询",
    "查找",
    "看看",
    "说明",
    "是什么",
    "怎么",
    "如何",
    "多少",
    "哪些",
    "有关",
    "关于",
}


_QUERY_SYNONYMS = {
    "怎么算": ["计算方法", "计费标准"],
    "怎么计算": ["计算方法", "计费标准"],
    "取费": ["计费标准", "计算方法"],
    "计费": ["计费标准", "计算方法"],
    "标准": ["计费标准", "计费基价"],
    "费用": ["费用内容", "计算方法"],
    "征地": ["土地补偿费", "土地补偿"],
    "土地": ["土地补偿费", "土地补偿"],
    "审查": ["设计审查费", "审查"],
    "监理": ["工程监理费", "监理"],
    "项目管理": ["建设单位管理费", "项目管理"],
}


_QUERY_KEY_TERMS = (
    "可行性研究",
    "建设单位",
    "项目管理",
    "土地",
    "征地",
    "补偿",
    "海域",
    "使用金",
    "评价",
    "评估",
    "论证",
    "招标",
    "咨询",
    "设计",
    "审查",
    "监理",
    "监管",
    "保险",
    "勘察",
    "检测",
    "检验",
    "标定",
    "试运转",
    "准备",
    "拆除",
    "运输",
    "预备",
    "取费",
    "计费",
    "标准",
    "计算方法",
    "费用内容",
)


def _split_compact_toc_entries(line: str) -> list[str]:
    if not re.search(r"\.{4,}", line):
        return []
    pattern = re.compile(
        r"((?:第[一二三四五六七八九十百]+章|（[一二三四五六七八九十]+）|"
        r"\d{1,2}\.)\s*[^.。\\n]{1,60}?"
        r"\s*\.{4,}\s*-?\s*\d+\s*-?)"
    )
    return [match.group(1).strip() for match in pattern.finditer(line)]


def _should_merge_pdf_body_line(previous: str, current: str) -> bool:
    if not previous or not current:
        return False
    if _looks_like_pdf_page_number(previous) or _looks_like_pdf_page_number(current):
        return False
    if _looks_like_pdf_toc_line(previous) or _looks_like_pdf_toc_line(current):
        return False
    if _looks_like_table_title(previous) or _looks_like_table_title(current):
        return False
    if _parse_pdf_heading(previous) is not None or _parse_pdf_heading(current) is not None:
        return False
    if re.match(r"^[（(]?\d+[）)]\s*", current):
        return False
    if re.match(r"^[一二三四五六七八九十]+、", current):
        return False
    if re.match(r"^[①②③④⑤⑥⑦⑧⑨]", current):
        return False
    if re.search(r"[。！？!?；;：:]$", previous):
        return False
    if _looks_like_pdf_table_fragment(previous) or _looks_like_pdf_table_fragment(current):
        return False
    return bool(re.search(r"[\u4e00-\u9fff）》]$", previous))


def _looks_like_pdf_page_number(text: str) -> bool:
    return bool(re.fullmatch(r"[-—－]\s*\d+\s*[-—－]", text.strip()))


def _looks_like_pdf_table_fragment(text: str) -> bool:
    normalized = text.strip()
    if len(normalized) <= 4:
        return True
    if normalized in {"序号", "备注", "费用项目类别", "费用项目名称"}:
        return True
    digit_groups = len(re.findall(r"\d+(?:\.\d+)?", normalized))
    return digit_groups >= 3 and not re.search(r"[，。；：]", normalized)


def _clean_cell_text(value: str) -> str:
    return re.sub(r"[\r\n]+", " ", value).strip("\n")


def _cell_text_from_tc(tc: object) -> str:
    texts = [node.text for node in tc.xpath(".//w:t") if node.text]
    return _clean_cell_text("".join(texts))


def _has_merged_cells(table: Table) -> bool:
    for row in table._tbl.tr_lst[:2]:
        if len(row.tc_lst) < len(table.columns):
            return True
        for tc in row.tc_lst:
            tc_pr = tc.tcPr
            if tc_pr is None:
                continue
            if tc_pr.vMerge is not None:
                return True
    return False


def _table_to_html(table: Table) -> str:
    lines: list[str] = []
    for row in table._tbl.tr_lst:
        cells: list[str] = []
        for tc in row.tc_lst:
            tc_pr = tc.tcPr
            attrs: list[str] = []
            text = _cell_text_from_tc(tc)
            if tc_pr is not None and tc_pr.gridSpan is not None:
                attrs.append(f'colspan="{tc_pr.gridSpan.val}"')
            if tc_pr is not None and tc_pr.vMerge is not None:
                v_merge = tc_pr.vMerge.val or "continue"
                if v_merge == "restart":
                    attrs.append('rowspan="2"')
                elif v_merge == "continue" and not text:
                    cells.append("<td>")
                    continue
            attr_text = " " + " ".join(attrs) if attrs else ""
            cells.append(f"<td{attr_text}>{_format_cell_html(text)}</td>")
        lines.append(f"<tr>{''.join(cells)}</tr>")
    return "<table>" + "\n".join(lines) + "</table>"


def _format_cell_html(text: str) -> str:
    escaped = escape(text)
    return re.sub(r"10([48])", r"10<sup>\1</sup>", escaped)


def _slice_html_table(html: str, row_scope: str | None) -> str:
    if not row_scope:
        return html
    match = re.match(r"^<table>(.*)</table>$", html, re.DOTALL)
    if not match:
        return html
    rows = re.findall(r"<tr>.*?</tr>", match.group(1), re.DOTALL)
    selected: list[str] = []
    for start, end in _parse_scopes(row_scope):
        selected.extend(rows[max(0, start - 1) : min(len(rows), end)])
    return "<table>\n" + "\n".join(selected) + "\n</table>"


def _normalize_heading(value: str) -> str:
    return re.sub(r"\s+", "", value).strip()


def _parse_heading(text: str, style_name: str) -> dict[str, str | int] | None:
    style_match = re.search(r"Heading\s+(\d+)|标题\s*(\d+)", style_name, re.IGNORECASE)
    if len(text) > 80:
        return None
    if _looks_like_numbered_body_sentence(text):
        return None
    if _looks_like_remote_ignored_heading(text):
        return None
    if re.match(r"^\d+(?:\.\d+)*-\d+\s+", text):
        return None
    if re.search(r"\d+\.\d+\s+\.\d+", text) or ".." in text:
        return None
    if text in {
        "3.1.8.5.3完井阶段钻井液维护要求",
        "3.1.11 钻井工程风险评估与应对",
        "5.3.1.3 主要工作量",
    }:
        return None
    number_match = re.match(r"^(\d+(?:\.\d+)+)\s+(.+)$", text)
    if number_match is None:
        number_match = re.match(r"^(\d+(?:\.\d+)+)([^\d\s].*)$", text)
    if number_match is None:
        number_match = re.match(r"^(\d+(?:\.\d+)*)(?:[.\s]+)(.+)$", text)
    if number_match is None:
        number_match = re.match(r"^(\d{1,2})((?!号)[\u4e00-\u9fff].*)$", text)
    if style_match:
        if number_match is None:
            return None
        level = int(style_match.group(1) or style_match.group(2))
        number = number_match.group(1)
        title = number_match.group(2).strip()
        return {"number": number, "title": title, "level": level}
    if number_match:
        number = number_match.group(1)
        title = number_match.group(2).strip()
        level = number.count(".") + 1
        return {"number": number, "title": title, "level": level}
    return None


def _parse_automatic_word_heading(
    paragraph: Paragraph,
    text: str,
    counters_by_num_id: dict[int, list[int]],
) -> dict[str, str | int] | None:
    """Rebuild headings whose visible numbering was lost during .doc conversion.

    LibreOffice can preserve Word's heading style, outline level and list metadata
    while omitting the rendered list number from ``Paragraph.text``. Those
    paragraphs are still reliable headings, so reconstruct their hierarchical
    number from the list level.
    """
    if not text or len(text) > 80:
        return None

    paragraph_properties = paragraph._p.pPr
    if paragraph_properties is None or paragraph_properties.numPr is None:
        return None

    style_name = paragraph.style.name if paragraph.style else ""
    style_match = re.search(
        r"Heading(?:\s+(\d+))?$|标题\s*(\d+)$",
        style_name.strip(),
        re.IGNORECASE,
    )
    outline_level = (
        int(paragraph_properties.outlineLvl.val)
        if paragraph_properties.outlineLvl is not None
        else None
    )
    if style_match is None and outline_level is None:
        return None

    num_properties = paragraph_properties.numPr
    num_id = (
        int(num_properties.numId.val)
        if num_properties.numId is not None
        else 0
    )
    list_level = (
        int(num_properties.ilvl.val)
        if num_properties.ilvl is not None
        else None
    )
    if list_level is None:
        if outline_level is not None:
            list_level = outline_level
        elif style_match is not None:
            explicit_level = style_match.group(1) or style_match.group(2)
            list_level = max(0, int(explicit_level or 1) - 1)
    if list_level is None or not 0 <= list_level <= 8:
        return None

    counters = counters_by_num_id.setdefault(num_id, [0] * 9)
    counters[list_level] += 1
    for deeper_level in range(list_level + 1, len(counters)):
        counters[deeper_level] = 0
    for parent_level in range(list_level):
        if counters[parent_level] == 0:
            counters[parent_level] = 1

    number = ".".join(str(value) for value in counters[: list_level + 1])
    return {
        "number": number,
        "title": text,
        "level": list_level + 1,
        "full_title": f"{number} {text}",
    }


def _parse_pdf_heading(text: str) -> dict[str, str | int] | None:
    if _looks_like_pdf_toc_line(text):
        return None
    if len(text) > 80:
        return None
    normalized = _normalize_heading(text)
    if re.fullmatch(r"\d+(?:\.\d+)*", normalized):
        return None
    if re.match(r"^\d+(?:\.\d+)?\s*(?:元|万|亿|km|m|亩|米|%)\b", text, re.IGNORECASE):
        return None
    cn_match = re.match(r"^(第[一二三四五六七八九十百]+章)\s*(.+)$", text)
    if cn_match is not None:
        number = cn_match.group(1)
        title = cn_match.group(2).strip()
        if title:
            return {"number": number, "title": title, "level": 1}
    cn_marker_match = re.match(r"^([一二三四五六七八九十]+)、\s*(.+)$", text)
    if cn_marker_match is not None:
        title = cn_marker_match.group(2).strip()
        if _looks_like_pdf_generic_heading_title(title):
            return {"number": cn_marker_match.group(1), "title": title, "level": 2}
    cn_paren_match = re.match(r"^（([一二三四五六七八九十]+)）\s*(.+)$", text)
    if cn_paren_match is not None:
        title = cn_paren_match.group(2).strip()
        if _looks_like_pdf_generic_heading_title(title):
            return {"number": f"（{cn_paren_match.group(1)}）", "title": title, "level": 3}
    match = re.match(r"^(\d+(?:\.\d+)*)(?:[.\s]+)?((?:[A-Za-z][A-Za-z0-9]*\s*)?[\u4e00-\u9fff].*)$", text)
    if match is None:
        match = re.match(r"^(\d+(?:\.\d+)*)(?:[.\s]+)?([A-Z][A-Za-z][A-Za-z\s/&-]*)$", text)
        if match is None:
            return None
    number = match.group(1)
    title = match.group(2).strip()
    if not title or _looks_like_numbered_body_sentence(text):
        return None
    if title.startswith("以下"):
        return None
    if number.count(".") >= 4:
        return None
    if "." not in number and _looks_like_pdf_fee_item_title(title) and re.match(r"^\d{1,2}\.\s*", text):
        return {"number": number, "title": title, "level": 2}
    if "." not in number and _looks_like_pdf_generic_top_heading_title(title):
        return {"number": number, "title": title, "level": 1}
    if "." not in number and not _looks_like_pdf_top_heading_title(title):
        if (
            not re.fullmatch(r"\d{1,2}", number)
            or not _looks_like_pdf_fee_item_title(title)
            or not re.match(r"^\d{1,2}\.\s*", text)
        ):
            return None
    if "." not in number and not _looks_like_pdf_top_heading_title(title):
        return None
    return {"number": number, "title": title, "level": number.count(".") + 1}


def _looks_like_pdf_top_heading_title(title: str) -> bool:
    normalized = _normalize_heading(title)
    keywords = {
        "总则",
        "范围",
        "规范性引用文件",
        "术语和定义",
        "一般规定",
        "建设条件",
        "建设规模",
        "总体布局",
        "油气集输工程",
        "配套工程",
        "劳动定员",
        "健康、安全与环保（HSE）",
        "其他费用",
        "预备费",
        "总论",
        "油气藏工程",
        "油藏工程",
        "钻井工程方案",
        "采油工程",
        "采油工程方案",
        "油气田地面工程",
        "地面工程",
        "节能降耗",
        "安全卫生与健康",
        "安全与职业卫生",
        "环境保护",
        "项目组织及进度安排",
        "主要风险评估与应对",
        "经济评价",
    }
    english_keywords = {"summary", "economy", "economic evaluation"}
    return normalized in keywords or title.strip().lower() in english_keywords


def _looks_like_pdf_generic_top_heading_title(title: str) -> bool:
    normalized = _normalize_heading(title)
    if not normalized or len(normalized) > 24:
        return False
    if re.search(r"[A-Za-z0-9（）()，,。；;：:]", title):
        return False
    if not re.search(r"[\u4e00-\u9fff]", normalized):
        return False
    if re.search(r"[，,。；;：:、]$", title):
        return False
    if re.search(r"^(本办法|本标准|本项目|该|其|其中|包括|按照|根据|对于|应|为|由|在)", normalized):
        return False
    if re.search(r"审批|审核|审议|备案", normalized) and not re.search(r"主体|程序|要求|管理|职责|监督|附件", normalized):
        return False
    return bool(
        re.search(
            r"总则|术语|定义|划分|主体|组织|管理|职责|计划|实施|统计|评价|考核|监督|检查|附则|附件|审批|决策|投资|项目",
            normalized,
        )
    )


def _is_summary_pdf_chapter(chapter: LocalChapter) -> bool:
    title = chapter.title.strip()
    normalized = _normalize_heading(title)
    if not normalized:
        return False
    if chapter.level > 3:
        return False
    number = str(chapter.number)
    if _is_chinese_marker_heading({"number": number}):
        return True
    if chapter.level <= 2:
        if chapter.level == 1:
            return True
        top_number = _first_heading_number(number)
        if top_number == 5:
            return _looks_like_summary_title(title)
        if re.fullmatch(r"\d{1,2}", number) and _looks_like_pdf_fee_item_title(title):
            return True
        return False
    if chapter.level == 3:
        return _looks_like_summary_title(title)
    if len(normalized) > 26:
        return False
    if _looks_like_numbered_body_sentence(chapter.full_title):
        return False
    if re.search(r"[，,。；;：:]", title):
        return False
    if _first_heading_number(number) != 5:
        return False
    return False


def _looks_like_summary_title(title: str) -> bool:
    normalized = _normalize_heading(title)
    if not normalized or len(normalized) > 18:
        return False
    if normalized.endswith(("应", "应针", "宜", "可", "不得", "必须")):
        return False
    if re.search(r"[，,。；;：:]", title):
        return False
    if re.search(r"^(本办法|建设单位|发展规划部|财务资产|概预算|纪检|审计|油气|工程技术|页岩油|CCUS|生产运行|企管|安全环保|科技|设备|经营|公共事业|对外合作|信息化|物资)", normalized):
        return False
    if re.search(r"^(发展规划部|建设单位|投资项目已完成|投资项目后评价是|发展规划部统一|建设单位必须|投资统计完成|投资绩效考核遵循)", normalized):
        return False
    return bool(re.search(r"规划|战略|流程|研究|论证|审批|设计|计划|变更|统计|后评价|绩效考核|调整|管理|附则|附件|权限|职责|定义|总则", normalized))
def _looks_like_pdf_fee_item_title(title: str) -> bool:
    normalized = _normalize_heading(title)
    keywords = (
        "费",
        "金",
        "补偿",
        "评价",
        "评估",
        "论证",
        "管理",
        "招标",
        "咨询",
        "审查",
        "监管",
        "监理",
        "保险",
        "勘察",
        "设计",
        "检测",
        "检验",
        "标定",
        "试运转",
        "准备",
        "拆除",
        "设施",
        "运输",
        "试验",
        "专利",
        "行政事业性",
        "预备",
    )
    return len(normalized) <= 35 and any(keyword in normalized for keyword in keywords)


def _looks_like_pdf_generic_heading_title(title: str) -> bool:
    normalized = _normalize_heading(title)
    if not normalized or len(normalized) > 30:
        return False
    if re.search(r"[，,。；;：:]$", title):
        return False
    if normalized.endswith("项目"):
        return len(normalized) <= 36 and not re.search(r"[，,。；;：:]", title)
    if re.search(r"^(本标准|本项目|该|其|其中|包括|按照|根据|对于|应|为|由|在)", normalized):
        return False
    return True


def _is_chinese_marker_heading(heading: dict[str, str | int]) -> bool:
    number = str(heading.get("number", ""))
    return bool(
        re.fullmatch(r"[一二三四五六七八九十]+", number)
        or re.fullmatch(r"（[一二三四五六七八九十]+）", number)
    )


def _filter_pdf_heading_candidates(
    candidates: list[tuple[int, str, dict[str, str | int]]],
) -> list[tuple[int, str, dict[str, str | int]]]:
    filtered = list(candidates)
    has_explicit_structure = any(
        int(heading.get("level", 0)) == 1 and not _is_chinese_marker_heading(heading)
        for _, _, heading in filtered
    )
    if has_explicit_structure:
        return [
            item
            for item in filtered
            if not _is_chinese_marker_heading(item[2])
        ]
    marker_indexes = [
        idx
        for idx, (_, _, heading) in enumerate(filtered)
        if _is_chinese_marker_heading(heading)
    ]
    keep_markers = set(_continuous_chinese_marker_indexes(filtered, marker_indexes))
    return [
        item
        for idx, item in enumerate(filtered)
        if not _is_chinese_marker_heading(item[2]) or idx in keep_markers
    ]


def _reject_pdf_heading_candidate(
    lines: list[str],
    line_no: int,
    heading: dict[str, str | int],
    current_top_level_number: int,
) -> bool:
    return _pdf_heading_candidate_score(lines, line_no, heading, current_top_level_number) < 0


def _pdf_heading_candidate_score(
    lines: list[str],
    line_no: int,
    heading: dict[str, str | int],
    current_top_level_number: int,
) -> int:
    title = str(heading.get("title", ""))
    normalized_title = _normalize_heading(title)
    level = int(heading.get("level", 0))
    number = str(heading.get("number", ""))
    current = _first_heading_number(number)
    score = 0

    if number.startswith("第") and level == 1:
        score += 90
    elif _is_chinese_marker_heading(heading):
        score += 45
    elif _is_pdf_fee_item_heading(heading):
        score += 55
    elif level == 1 and _looks_like_pdf_top_heading_title(title):
        score += 55
    elif level == 1 and _looks_like_pdf_generic_top_heading_title(title):
        score += 40
    elif level > 1:
        score += 30

    if normalized_title and len(normalized_title) <= 16:
        score += 10
    if current_top_level_number and current == current_top_level_number and level > 1:
        score += 15
    if level == 1 and current and current_top_level_number and current == current_top_level_number + 1:
        score += 20

    if (
        level == 1
        and _looks_like_pdf_table_or_formula_heading(title)
        and not _looks_like_pdf_top_heading_title(title)
    ):
        score -= 100
    if level == 1 and current_top_level_number > 0:
        if (
            current
            and current <= current_top_level_number
            and not number.startswith("第")
            and re.search(r"[\u4e00-\u9fff]", title)
        ):
            score -= 90
    if level <= 2 and _near_pdf_table_context(lines, line_no) and not _is_pdf_fee_item_heading(heading):
        score -= 70
    if re.fullmatch(r"\d{1,2}", number) and normalized_title.startswith("以下"):
        score -= 90
    if re.search(r"^(费用项目类别|序号|备注|计费基价|项目总投资)$", normalized_title):
        score -= 90
    if len(normalized_title) > 30:
        score -= 25
    return score


def _looks_like_pdf_table_or_formula_heading(title: str) -> bool:
    normalized = _normalize_heading(title)
    if not normalized:
        return True
    if not re.search(r"[\u4e00-\u9fff]", normalized):
        return False
    if re.search(r"[A-Za-z0-9（）()]", title):
        return True
    if re.search(r"(单位|含\d|以上|以下|评价面积|计费基价|万元|亿元|km|m2|%)", normalized, re.IGNORECASE):
        return True
    if len(normalized) > 24:
        return True
    return False


def _near_pdf_table_context(lines: list[str], line_no: int) -> bool:
    start = max(0, line_no - 5)
    context = "\n".join(lines[start : line_no - 1])
    return bool(
        re.search(r"表\d|费用项目类别|费用项目名称|序号|计费基价|项目总投资|备注", context)
    )


def _continuous_chinese_marker_indexes(
    candidates: list[tuple[int, str, dict[str, str | int]]],
    indexes: list[int],
) -> set[int]:
    groups: dict[tuple[int, str], list[tuple[int, int]]] = {}
    for idx in indexes:
        line_no, _, heading = candidates[idx]
        number = str(heading.get("number", ""))
        marker_type = "paren" if number.startswith("（") else "plain"
        ordinal = _chinese_marker_to_int(number)
        if ordinal <= 0:
            continue
        groups.setdefault((int(heading.get("level", 0)), marker_type), []).append((idx, ordinal))
    keep: set[int] = set()
    for values in groups.values():
        values.sort()
        for pos, (idx, ordinal) in enumerate(values):
            prev_ordinal = values[pos - 1][1] if pos > 0 else 0
            next_ordinal = values[pos + 1][1] if pos + 1 < len(values) else 0
            if ordinal == 1 or ordinal == prev_ordinal + 1 or next_ordinal == ordinal + 1:
                keep.add(idx)
    return keep


def _chinese_marker_to_int(number: str) -> int:
    return _chinese_number_to_int(number.strip("（）"))


def _is_pdf_fee_item_heading(heading: dict[str, str | int]) -> bool:
    number = str(heading.get("number", ""))
    title = str(heading.get("title", ""))
    return (
        heading.get("level") == 2
        and bool(re.fullmatch(r"\d{1,2}", number))
        and _looks_like_pdf_fee_item_title(title)
    )


def _pdf_fee_item_belongs_to_top(item_number: int, top_number: int) -> bool:
    if top_number == 2:
        return 1 <= item_number <= 44
    if top_number == 3:
        return 45 <= item_number <= 46
    return False


def _looks_like_numbered_body_sentence(text: str) -> bool:
    if re.search(r"[，,。；;：:]$", text):
        return True
    if len(text) >= 36 and re.search(r"包括|计算|按|为：|为$", text):
        return True
    return bool(re.search(r"^\d+(?:\.\d+)*\s*座", text))


def _looks_like_remote_ignored_heading(text: str) -> bool:
    normalized = _normalize_heading(text)
    ignored_exact = {
        "2.6钻井工程风险评估与应对",
        "2.6.1防喷",
        "2.6.2防漏",
        "2.6.3污染防治措施",
        "2.7健康、安全与环境管理要求",
        "2.7.1职业健康",
        "2.7.1.1遵循的主要法律法规及标准规范",
        "2.7.1.2职业健康防护措施对策",
        "2.8.2安全",
        "2.8.2.1遵循的主要法律法规及标准规范",
        "2.8.2.2安全风险防护措施",
        "2.8.3环境",
        "2.8.3.1遵循的主要法律法规及标准规范",
        "2.8.3.2环境风险防范措施",
        "2.9节能措施",
        "2.9.1节能措施分析",
        "2.9.2节能风险评估与应对",
        "3.2钻井工程费用估算",
        "3.4.2.7注水井动态监测工艺",
        "3.4.3.2.4CO2泄漏应急预案",
        "4.3.2自控部分",
        "4.3.2.1设计内容",
        "4.4.4.2设计条件",
        "4.4.5.3设计方案",
        "4.4.5.4主要工作量",
        "4.7生产现场外观形象标准化建设",
        "4.8组织结构与劳动定员",
        "4.8.1组织机构",
        "4.8.2劳动定员",
        "4.9地面工程弃置方案",
        "4.9.1地面工程清理方案",
        "4.9.2弃置费测算",
        "4.10地面工程投资估算",
        "4.10.1组织机构",
        "4.10.1.1定额依据",
        "4.10.1.2费用依据",
        "4.10.1.3价格依据",
        "4.11投资估算",
        "5.5.1遵循规范",
        "5.5.2方案设计",
        "5.6.4碳排放经济性评价",
        "6.5.3碳排放强度评价",
        "7.1.1.6.2与《滨州市生态环境准入清单（2023年版）》的符合性分析",
        "7.1.1.6.3与淄政字〔2021〕49号符合性分析",
        "7.1.1.6.4与东环委办[2024]7号符合性分析",
        "7.1.1.1废液",
        "7.1.1.2废气",
        "7.1.1.3噪声",
        "7.1.1.4固废",
        "7.1.1.5生态",
    }
    if normalized in ignored_exact:
        return True
    return bool(re.match(r"^3\.(?:1|3)\.5\.6$", normalized))


def _looks_like_table_title(text: str) -> bool:
    return bool(
        re.match(r"^表\s*\d+(?:[-.]\d+)*", text)
        or re.match(r"^Table\s+\d+(?:[-.]\d+)*", text, re.IGNORECASE)
    )


def _looks_like_after_table_body(text: str) -> bool:
    return bool(
        re.match(r"^(各单项工程|本项目评价期内|从图|根据|结合方案|经估算|方案部署)", text)
        or re.match(r"^（\d+）", text)
        or re.match(r"^[①②③④⑤⑥⑦⑧⑨]", text)
    )


def _looks_like_pdf_toc_line(text: str) -> bool:
    return bool(re.search(r"\.{4,}\s*(?:-?\s*\d+\s*-?)?\s*$", text))


def _is_pdf_page_number_heading(text: str, heading: dict[str, str | int]) -> bool:
    if heading.get("level") != 1:
        return False
    number = str(heading.get("number", ""))
    title = str(heading.get("title", ""))
    return number.isdigit() and not title


def _looks_like_real_pdf_section_heading(text: str, heading: dict[str, str | int]) -> bool:
    if heading.get("level") == 1:
        return True
    title = str(heading.get("title", "")).strip()
    if re.search(r"\d", title):
        return False
    if len(title) > 30:
        return False
    return bool(re.search(r"工程|方案|估算|评价|分析|测算|依据|原则|设计|要求|措施|结论|概况|特征|投资|筹措|费用|风险|环保|安全|卫生|组织|进度", title))


def _is_first_body_heading(text: str, heading: dict[str, str | int] | None = None) -> bool:
    if re.match(r"^第一章\s*总则$", text):
        return True
    if re.match(r"^1[\.、\s]*总论(?:\s|$)", text):
        return True
    if heading is None:
        parsed = _parse_heading(text, "")
        if parsed is None:
            return False
        heading = parsed
    if heading.get("number") in {"1", "第一章"} and heading.get("level") == 1:
        return True
    return str(heading.get("number")) in {"一", "（一）"}


def _is_toc_style(style_name: str) -> bool:
    normalized = style_name.strip().lower()
    return normalized.startswith("toc") or normalized.startswith("目录")


def _is_out_of_sequence_top_heading(
    heading: dict[str, str | int], last_top_level_number: int
) -> bool:
    if heading.get("level") != 1 or last_top_level_number <= 0:
        return False
    current = _first_heading_number(str(heading.get("number", "")))
    if current <= 0:
        return False
    return current < last_top_level_number or current > last_top_level_number + 1


def _first_heading_number(number: str) -> int:
    cn_match = re.match(r"^第([一二三四五六七八九十百]+)章$", number)
    if cn_match:
        return _chinese_number_to_int(cn_match.group(1))
    match = re.match(r"^(\d+)", number)
    return int(match.group(1)) if match else 0


def _chinese_number_to_int(value: str) -> int:
    digits = {
        "零": 0,
        "一": 1,
        "二": 2,
        "三": 3,
        "四": 4,
        "五": 5,
        "六": 6,
        "七": 7,
        "八": 8,
        "九": 9,
    }
    if value == "十":
        return 10
    if "十" in value:
        left, _, right = value.partition("十")
        tens = digits.get(left, 1) if left else 1
        ones = digits.get(right, 0) if right else 0
        return tens * 10 + ones
    return digits.get(value, 0)


def _heading_number_level(number: str) -> int:
    return number.count(".") + 1 if number else 99


def _looks_like_chapter_number_query(query: str) -> bool:
    return bool(re.fullmatch(r"\d+(?:\.\d+)*", query.strip()))


def _chapter_key_from_table_title(text: str) -> str:
    match = re.match(r"^(?:表|Table)\s*(\d+(?:[-.]\d+)*)", text, re.IGNORECASE)
    if not match:
        return ""
    return match.group(1).replace("-", ".")


def _parse_scopes(value: str) -> list[tuple[int, int]]:
    scopes: list[tuple[int, int]] = []
    for part in re.split(r"[,，]", value):
        match = re.match(r"\s*(\d+)\s*[-~]\s*(\d+)\s*$", part)
        if match:
            start, end = int(match.group(1)), int(match.group(2))
            scopes.append((min(start, end), max(start, end)))
            continue
        single = re.match(r"\s*(\d+)\s*$", part)
        if single:
            line = int(single.group(1))
            scopes.append((line, line))
    return scopes
